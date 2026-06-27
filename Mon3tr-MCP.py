"""
Mon3tr-MCP — 明日方舟 MCP 工具集
支持 Windows / macOS / Linux（含 Android Termux）

工具列表:
  搜索/抓取: bing_search · fetch_page · fetch_prts_wiki
  游戏数据:  fetch_gamedata（按仓库相对路径从 GitHub 拉取，带缓存）
  地图解析:  parse_map · get_cell_info · get_map_legend
  敌人数据:  get_level_enemies · get_enemy_by_id
  文档工具:  generate_operator_profile · edit_document · edit_excel · edit_pdf
  文件工具:  edit_txt · edit_json
  干员查询:  query_operator
  剧情工具:  query_story
  CSGO赛事:  csgo_matches · csgo_rankings · csgo_tournaments · csgo_player_stats
             数据源: HLTV(抓取) + Liquipedia(API) + 5EPlay(抓取)
  注: parse_map / get_level_enemies / get_enemy_by_id 均支持本地路径、完整 URL
      或仓库相对路径（自动补全远端地址）
"""

from __future__ import annotations

import json
import os
import re
import sys
from pathlib import Path
from typing import Optional

import requests
from bs4 import BeautifulSoup
from mcp.server.fastmcp import FastMCP

# ── 远端游戏数据源 ────────────────────────────────────────────────
_GAMEDATA_BASE_URL = (
    "https://raw.githubusercontent.com/Kengxxiao/ArknightsGameData/master/zh_CN"
)
_json_cache: dict = {}

# ── Windows 控制台 UTF-8 输出修正 ────────────────────────────────
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

mcp = FastMCP("Mon3tr-MCP", host="127.0.0.1", port=8000)


# ══════════════════════════════════════════════════════════════════
# 搜索 / 网页抓取工具
# ══════════════════════════════════════════════════════════════════

_MOBILE_UA = (
    "Mozilla/5.0 (Linux; Android 14; Pixel 8) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/120.0.0.0 Mobile Safari/537.36"
)


@mcp.tool()
def bing_search(query: str, num: int = 10, domain: str = "cn.bing.com") -> str:
    """
    用 Bing 搜索网页，返回标题、链接和摘要。

    参数:
        query: 搜索关键词
        num: 返回结果数量（默认10）
        domain: Bing 域名，可选 "cn.bing.com"（国内版，默认）或 "www.bing.com"（国际版）
    """
    if domain not in ("cn.bing.com", "www.bing.com"):
        domain = "cn.bing.com"
    headers = {
        "User-Agent": _MOBILE_UA,
        "Accept-Language": "zh-CN,zh;q=0.9",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Referer": f"https://{domain}/",
    }
    url = f"https://{domain}/search?q={query}&count={num}"
    res = requests.get(url, headers=headers, timeout=10)
    soup = BeautifulSoup(res.text, "html.parser")
    results = []
    for li in soup.select("li.b_algo"):
        title = li.select_one("h2")
        link = li.select_one("a")
        snippet = li.select_one(".b_caption p")
        if title and link:
            results.append(
                f"标题：{title.get_text()}\n"
                f"链接：{link['href']}\n"
                f"摘要：{snippet.get_text() if snippet else '无'}"
            )
    return "\n---\n".join(results) if results else "没有找到结果"


@mcp.tool()
def fetch_page(url: str) -> str:
    """抓取指定网页的正文内容"""
    headers = {
        "User-Agent": _MOBILE_UA,
        "Accept-Language": "zh-CN,zh;q=0.9",
    }
    try:
        res = requests.get(url, headers=headers, timeout=10)
        res.encoding = res.apparent_encoding
        soup = BeautifulSoup(res.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        content = "\n".join(lines)
        return content[:3000] + "..." if len(content) > 3000 else content
    except Exception as e:
        return f"抓取失败：{e}"


@mcp.tool()
def fetch_prts_wiki(character_name: str) -> str:
    """从 prts.wiki 抓取明日方舟角色详细信息"""
    headers = {
        "User-Agent": _MOBILE_UA,
        "Accept-Language": "zh-CN,zh;q=0.9",
    }
    url = f"https://prts.wiki/w/{character_name}"
    try:
        res = requests.get(url, headers=headers, timeout=10)
        res.encoding = "utf-8"
        soup = BeautifulSoup(res.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "table"]):
            tag.decompose()
        content = soup.find("div", {"id": "mw-content-text"})
        if not content:
            return "未找到该角色页面，请检查角色名称"
        text = content.get_text(separator="\n")
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        result = "\n".join(lines)
        return result[:5000] + "..." if len(result) > 5000 else result
    except Exception as e:
        return f"抓取失败：{e}"


# ══════════════════════════════════════════════════════════════════
# 通用 JSON 加载（本地路径 / HTTP URL，带内存缓存）
# ══════════════════════════════════════════════════════════════════

def _load_json(path_or_url: str) -> dict:
    if path_or_url.startswith(("http://", "https://")):
        if path_or_url not in _json_cache:
            res = requests.get(path_or_url, timeout=15)
            res.raise_for_status()
            _json_cache[path_or_url] = res.json()
        return _json_cache[path_or_url]
    with Path(path_or_url).open("r", encoding="utf-8") as f:
        return json.load(f)


def _is_accessible(path_or_url: str) -> bool:
    return path_or_url.startswith(("http://", "https://")) or Path(path_or_url).is_file()


# ══════════════════════════════════════════════════════════════════
# 地图解析辅助函数
# ══════════════════════════════════════════════════════════════════

_TILE_MAP = {
    "tile_forbidden":        "禁",
    "tile_wall":             "高",
    "tile_road":             "路",
    "tile_floor":            "不",
    "tile_hole":             "穴",
    "tile_telin":            "进",
    "tile_telout":           "出",
    "tile_fence_bound":      "围",
    "tile_start":            "红",
    "tile_end":              "蓝",
    "tile_mpprts_enemy_born": "红",
}

_FLY_START_TILES = {"tile_flystart"}

_CELL_DESC = {
    "禁": "禁入区 - 不可通行/不可部署",
    "高": "高台 - 可部署高台干员",
    "路": "可部署位 - 可部署地面干员",
    "不": "不可部署位 - 可通行但不可部署",
    "穴": "地穴",
    "装": "装置 - 预设装置/陷阱",
    "蓝": "目标点 - 防守点(蓝门)",
    "红": "侵入点 - 敌人生成点(红门)",
    "进": "传送进入点",
    "出": "传送离开点",
    "围": "围墙 - 可部署围墙",
    "飞": "飞行起点 - 飞行单位生成点",
    "?": "未知区域",
}

_LEVEL_TYPE_NAME = {0: "普通", 1: "精英", 2: "BOSS"}


def _parse_arknights_map_data(data: dict):
    """解析已加载的关卡 JSON dict，返回 (game_grid, blue_doors, red_doors, fly_doors, data)"""
    map_matrix = data["mapData"]["map"]
    tiles = data["mapData"]["tiles"]
    rows = len(map_matrix)
    cols = len(map_matrix[0])

    grid = [["?" for _ in range(cols)] for _ in range(rows)]
    fly_start_points: set = set()

    for y in range(rows):
        for x in range(cols):
            tile_index = map_matrix[y][x]
            tile_key = tiles[tile_index].get("tileKey") if tile_index < len(tiles) else None
            if tile_key in _TILE_MAP:
                grid[rows - y - 1][x] = _TILE_MAP[tile_key]
            elif tile_key in _FLY_START_TILES:
                fly_start_points.add((y, x))
            # tile_start / tile_end 留作 '?' 占位，后续被门覆盖

    # 标记预设装置（在标记门之前，确保门的排除条件能正确检查 '装'）
    if "predefines" in data and "tokenInsts" in data["predefines"]:
        for token in data["predefines"]["tokenInsts"]:
            r = token["position"]["row"]
            c = token["position"]["col"]
            grid[r][c] = "装"

    blue_doors: set = set()
    for route in data.get("routes", []):
        sp, ep = route["startPosition"], route["endPosition"]
        if sp["row"] != ep["row"] or sp["col"] != ep["col"]:
            er, ec = ep["row"], ep["col"]
            if grid[er][ec] not in {"穴", "装"}:
                grid[er][ec] = "蓝"
                blue_doors.add((er, ec))

    red_doors: set = set()
    for route in data.get("routes", []):
        sp, ep = route["startPosition"], route["endPosition"]
        if sp["row"] != ep["row"] or sp["col"] != ep["col"]:
            sr, sc = sp["row"], sp["col"]
            if grid[sr][sc] not in {"路", "不", "围", "装", "穴"}:
                grid[sr][sc] = "红"
                red_doors.add((sr, sc))

    # 标记飞行单位生成点（覆盖红门）
    for fy, fx in fly_start_points:
        grid[rows - fy - 1][fx] = "飞"

    # 坐标系翻转：game_grid[0] 对应屏幕上方
    game_grid = [["?" for _ in range(cols)] for _ in range(rows)]
    for row in range(rows):
        for col in range(cols):
            game_grid[rows - 1 - row][col] = grid[row][col]

    # fly_start_points 是 JSON 坐标，转为与 blue_doors/red_doors 一致的格式
    fly_doors = {(rows - 1 - fy, fx) for fy, fx in fly_start_points}

    return game_grid, blue_doors, red_doors, fly_doors, data


def _build_text_map(game_grid, blue_doors, red_doors) -> str:
    rows = len(game_grid)
    cols = len(game_grid[0])
    lines = ["左上角(0,0)"]
    for y in range(rows):
        row_str = f"    y={y} [" + "][".join(game_grid[y]) + "]"
        for r, c in red_doors:
            if rows - 1 - r == y:
                row_str += f" <- 红门({c},{rows-1-r})"
        lines.append(row_str)
    x_width = max(len(str(cols - 1)), 2)
    x_axis = "        " + "".join(
        f"x={x}".ljust(x_width + 3) for x in range(cols)
    )
    lines.append(x_axis)
    lines.append("\n图例:")
    lines.append(
        "禁-禁入区 高-高台 路-可部署位 不-不可部署位 穴-地穴 "
        "装-装置 蓝-防守点 红-侵入点 进-传送进入 出-传送离开 围-围墙 飞-飞行起点"
    )
    return "\n".join(lines)


def _get_cell_description(cell_type: str) -> str:
    return _CELL_DESC.get(cell_type, "未知区域")


# ══════════════════════════════════════════════════════════════════
# 地图 MCP 工具
# ══════════════════════════════════════════════════════════════════

@mcp.tool()
def parse_map(json_file_path: str) -> str:
    """
    解析明日方舟关卡 JSON 并返回文字地图及关卡基本信息。支持本地路径或 URL。

    参数:
        json_file_path: 关卡 JSON 的本地路径或完整 URL。
                        也可传入仓库相对路径（如 gamedata/levels/obt/main/level_main_01-08.json），
                        将自动拼接远端地址。

    返回:
        文字地图 + 红/蓝门坐标 + 关卡选项信息
    """
    if not json_file_path.startswith(("http://", "https://")):
        if not Path(json_file_path).is_file():
            # 尝试视作仓库相对路径从远端获取
            json_file_path = f"{_GAMEDATA_BASE_URL}/{json_file_path.lstrip('/')}"
    try:
        data = _load_json(json_file_path)
        game_grid, blue_doors, red_doors, fly_doors, data = _parse_arknights_map_data(data)
    except requests.HTTPError as e:
        return f"错误: 远端请求失败 {e}"
    except json.JSONDecodeError:
        return f"错误: 不是有效的 JSON"
    except KeyError as e:
        return f"错误: JSON 缺少必要字段 {e}"
    except Exception as e:
        return f"处理时发生错误: {e}"

    rows = len(game_grid)
    text_map = _build_text_map(game_grid, blue_doors, red_doors)
    blue_list = ", ".join(f"({c},{rows-1-r})" for r, c in sorted(blue_doors)) or "无"
    red_list  = ", ".join(f"({c},{rows-1-r})" for r, c in sorted(red_doors))  or "无"
    fly_list  = ", ".join(f"({c},{rows-1-r})" for r, c in sorted(fly_doors)) or "无"
    opts = data.get("options", {})
    info_lines = [
        f"地图尺寸: {rows} 行 x {len(game_grid[0])} 列",
        f"部署上限: {opts.get('characterLimit', '未知')} 个干员",
        f"生命值: {opts.get('maxLifePoint', '未知')}",
        f"初始费用: {opts.get('initialCost', '未知')}",
        f"蓝门坐标: {blue_list}",
        f"红门坐标: {red_list}",
        f"飞行起点坐标: {fly_list}",
    ]
    return text_map + "\n\n" + "\n".join(info_lines)


@mcp.tool()
def get_cell_info(cell_type: str) -> str:
    """
    查询地图格子类型的含义。

    参数:
        cell_type: 单字符类型，如 '高'、'路'、'红' 等
    """
    return _get_cell_description(cell_type)


@mcp.tool()
def get_map_legend() -> str:
    """返回完整的地图图例说明。"""
    return "\n".join(f"{k}: {v}" for k, v in _CELL_DESC.items())


# ══════════════════════════════════════════════════════════════════
# 敌人数据辅助函数
# ══════════════════════════════════════════════════════════════════

def _load_enemy_db(db_path_or_url: str) -> dict:
    raw = _load_json(db_path_or_url)
    return {entry["key"]: entry["value"] for entry in raw["enemies"]}


def _mv(field):
    """取 m_value（当 m_defined=True 时），否则返回 None。"""
    if field and field.get("m_defined"):
        return field.get("m_value")
    return None


def _merge_enemy_data(base: dict, overwrite: Optional[dict]) -> dict:
    """将 overwrittenData 中 m_defined=True 的字段覆盖到 base 上。"""
    if not overwrite:
        return base
    merged = dict(base)
    for field in (
        "name", "description", "prefabKey", "applyWay", "motion",
        "enemyTags", "lifePointReduce", "levelType", "rangeRadius",
        "numOfExtraDrops", "viewRadius", "notCountInTotal",
    ):
        ow_field = overwrite.get(field)
        if ow_field and ow_field.get("m_defined"):
            merged[field] = ow_field
    if "attributes" in overwrite:
        merged_attrs = dict(base.get("attributes", {}))
        for attr_key, attr_val in overwrite["attributes"].items():
            if attr_val and attr_val.get("m_defined"):
                merged_attrs[attr_key] = attr_val
        merged["attributes"] = merged_attrs
    return merged


def _format_enemy(enemy_id: str, enemy_data: dict, data_level: int = 0) -> str:
    attrs = enemy_data.get("attributes", {})

    def av(key):
        f = attrs.get(key)
        return _mv(f) if f else None

    lines = ["[{}]".format(enemy_id)]

    name = _mv(enemy_data.get("name"))
    if name:
        lines.append(f"名称: {name}")
    lines.append(f"数据级别: {data_level}")

    desc = _mv(enemy_data.get("description"))
    if desc:
        lines.append(f"描述: {re.sub(r'<[^>]+>', '', desc)}")

    tags = _mv(enemy_data.get("enemyTags"))
    if tags:
        lines.append(f"标签: {', '.join(tags)}")

    level_type = _mv(enemy_data.get("levelType"))
    if level_type is not None:
        lines.append(f"等级类型: {_LEVEL_TYPE_NAME.get(level_type, str(level_type))}")

    lp = _mv(enemy_data.get("lifePointReduce"))
    if lp is not None:
        lines.append(f"占用生命点: {lp}")

    rr = _mv(enemy_data.get("rangeRadius"))
    if rr:
        lines.append(f"攻击半径: {rr}")

    stat_fields = [
        ("maxHp",            "生命值"),
        ("atk",              "攻击力"),
        ("def",              "防御力"),
        ("magicResistance",  "法抗(%)"),
        ("moveSpeed",        "移动速度"),
        ("attackSpeed",      "攻击速度"),
        ("baseAttackTime",   "攻击间隔(s)"),
        ("massLevel",        "重量等级"),
        ("blockCnt",         "阻挡数"),
        ("hpRecoveryPerSec", "每秒回血"),
    ]
    stat_lines = [
        f"  {label}: {av(key)}"
        for key, label in stat_fields
        if av(key) not in (None, 0, 0.0)
    ]
    if stat_lines:
        lines.append("属性:")
        lines.extend(stat_lines)

    immune_keys = [
        ("stunImmune",     "眩晕免疫"),
        ("silenceImmune",  "沉默免疫"),
        ("sleepImmune",    "睡眠免疫"),
        ("frozenImmune",   "冻结免疫"),
        ("levitateImmune", "浮空免疫"),
    ]
    immunities = [label for key, label in immune_keys if av(key)]
    if immunities:
        lines.append(f"免疫: {', '.join(immunities)}")

    skills = enemy_data.get("skills")
    if skills:
        lines.append("技能:")
        for sk in skills:
            bb_str = ", ".join(
                f"{b['key']}={b['value']}"
                for b in sk.get("blackboard", [])
                if b.get("value") is not None
            )
            sk_line = f"  - {sk.get('prefabKey', '?')} (冷却:{sk.get('cooldown', 0)}s)"
            if bb_str:
                sk_line += f" [{bb_str}]"
            lines.append(sk_line)

    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════
# 敌人 MCP 工具
# ══════════════════════════════════════════════════════════════════

@mcp.tool()
def get_level_enemies(level_json_path: str, enemy_db_path: str) -> str:
    """
    读取关卡 JSON 的 enemyDbRefs，从 enemy_database.json 中查找敌人数据。
    两个参数均支持本地路径或 URL；也可传入仓库相对路径自动补全远端地址。

    参数:
        level_json_path: 关卡 JSON 路径/URL（如 gamedata/levels/obt/main/level_main_01-08.json）
        enemy_db_path:   enemy_database.json 路径/URL（如 gamedata/levels/enemydata/enemy_database.json）
    """
    def _resolve(p: str) -> str:
        if p.startswith(("http://", "https://")) or Path(p).is_file():
            return p
        return f"{_GAMEDATA_BASE_URL}/{p.lstrip('/')}"

    level_json_path = _resolve(level_json_path)
    enemy_db_path   = _resolve(enemy_db_path)
    try:
        level_data = _load_json(level_json_path)
        enemy_db   = _load_enemy_db(enemy_db_path)
    except requests.HTTPError as e:
        return f"错误: 远端请求失败 {e}"
    except json.JSONDecodeError as e:
        return f"错误: JSON 解析失败 - {e}"
    except Exception as e:
        return f"错误: {e}"

    refs = level_data.get("enemyDbRefs", [])
    if not refs:
        return "该关卡没有 enemyDbRefs 数据"

    results, not_found = [], []
    for ref in refs:
        eid   = ref["id"]
        level = ref.get("level", 0)
        ow    = ref.get("overwrittenData")
        if eid not in enemy_db:
            not_found.append(eid)
            continue
        entries    = enemy_db[eid]
        base_entry = next((e for e in entries if e["level"] == 0), entries[0])
        base_data  = base_entry["enemyData"]
        if level != 0:
            diff_entry = next((e for e in entries if e["level"] == level), None)
            if diff_entry:
                base_data = _merge_enemy_data(base_data, diff_entry["enemyData"])
        merged = _merge_enemy_data(base_data, ow)
        results.append(_format_enemy(eid, merged, data_level=level))

    output = f"关卡敌人数据（共 {len(results)} 种）\n" + "=" * 50 + "\n"
    output += "\n\n".join(results)
    if not_found:
        output += f"\n\n未在数据库中找到的敌人: {', '.join(not_found)}"
    return output


@mcp.tool()
def get_enemy_by_id(enemy_id: str, enemy_db_path: str, level: int = 0) -> str:
    """
    从 enemy_database.json 按 ID 查询单个敌人的数据。支持本地路径或 URL。

    参数:
        enemy_id:      敌人 ID，如 enemy_1000_gopro
        enemy_db_path: enemy_database.json 的路径/URL，或仓库相对路径
                       （如 gamedata/levels/enemydata/enemy_database.json）
        level:         数据等级（默认 0；精英/BOSS 关卡可能有 level 1+）
    """
    if not enemy_db_path.startswith(("http://", "https://")) and not Path(enemy_db_path).is_file():
        enemy_db_path = f"{_GAMEDATA_BASE_URL}/{enemy_db_path.lstrip('/')}"
    try:
        enemy_db = _load_enemy_db(enemy_db_path)
    except requests.HTTPError as e:
        return f"错误: 远端请求失败 {e}"
    except Exception as e:
        return f"错误: {e}"

    if enemy_id not in enemy_db:
        return f"未找到敌人: {enemy_id}"

    entries    = enemy_db[enemy_id]
    base_entry = next((e for e in entries if e["level"] == 0), entries[0])
    base_data  = base_entry["enemyData"]
    if level != 0:
        diff_entry = next((e for e in entries if e["level"] == level), None)
        if diff_entry:
            base_data = _merge_enemy_data(base_data, diff_entry["enemyData"])
    return _format_enemy(enemy_id, base_data, data_level=level)


# ══════════════════════════════════════════════════════════════════
# 远端游戏数据获取工具
# ══════════════════════════════════════════════════════════════════

@mcp.tool()
def fetch_gamedata(relative_path: str, max_chars: int = 3000) -> str:
    """
    从 GitHub 上的 ArknightsGameData 仓库（Kengxxiao/ArknightsGameData，zh_CN 分支）
    按相对路径获取游戏数据文件，结果以 JSON 字符串返回（带缓存）。

    参数:
        relative_path: 仓库内相对路径，如
                       gamedata/excel/character_table.json
                       gamedata/excel/skill_table.json
                       gamedata/levels/obt/main/level_main_01-08.json
                       gamedata/levels/enemydata/enemy_database.json
        max_chars:     返回内容最大字符数（默认 3000，大文件可调大）

    返回:
        JSON 字符串（超出 max_chars 时截断并提示）
    """
    url = f"{_GAMEDATA_BASE_URL}/{relative_path.lstrip('/')}"
    try:
        data = _load_json(url)
    except requests.HTTPError as e:
        return f"错误: 远端请求失败 {e}\n请求地址: {url}"
    except Exception as e:
        return f"错误: {e}"
    content = json.dumps(data, ensure_ascii=False, indent=2)
    if len(content) > max_chars:
        return content[:max_chars] + f"\n\n...(已截断，完整长度 {len(content)} 字符，请增大 max_chars)"
    return content


# ══════════════════════════════════════════════════════════════════
# 关卡查询工具
# ══════════════════════════════════════════════════════════════════

_STAGE_OVERVIEW_URL = (
    "https://raw.githubusercontent.com/MaaAssistantArknights/"
    "MaaAssistantArknights/dev-v2/resource/Arknights-Tile-Pos/overview.json"
)
_stage_overview_cache: Optional[dict] = None


def _load_stage_overview() -> dict:
    global _stage_overview_cache
    if _stage_overview_cache is None:
        res = requests.get(_STAGE_OVERVIEW_URL, timeout=15)
        res.raise_for_status()
        _stage_overview_cache = res.json()
    return _stage_overview_cache


@mcp.tool()
def lookup_stage(query: str) -> str:
    """
    通过关卡代号、名称或 stageId 查询关卡信息及游戏数据文件路径。

    参数:
        query: 关卡代号（如 "1-7"、"CE-5"、"GT-1"）、
               关卡名称（如 "当务之急"、"龙门外环"）或
               stageId（如 "main_01-07"）

    返回:
        匹配的关卡列表，包含代号、名称、stageId、
        游戏数据 JSON 路径（可用于 parse_map / fetch_gamedata）
    """
    try:
        overview = _load_stage_overview()
    except Exception as e:
        return f"错误: 无法获取关卡数据表 {e}"

    query_lower = query.lower().strip()
    results = []
    for entry in overview.values():
        code = entry.get("code", "")
        name = entry.get("name", "")
        stage_id = entry.get("stageId", "")
        if (query_lower in code.lower()
                or query_lower in name
                or query_lower in stage_id.lower()):
            results.append(entry)

    if not results:
        return f"未找到匹配 '{query}' 的关卡"

    lines = [f"找到 {len(results)} 个匹配关卡:", ""]
    for e in results[:20]:
        level_id = e.get("levelId", "")
        json_path = f"gamedata/levels/{level_id}.json"
        lines.append(f"  关卡代号: {e.get('code', '?')}")
        lines.append(f"  名称:     {e.get('name', '?')}")
        lines.append(f"  stageId:  {e.get('stageId', '?')}")
        lines.append(f"  地图尺寸: {e.get('width', '?')} x {e.get('height', '?')}")
        lines.append(f"  数据路径: {json_path}")
        lines.append("")
    if len(results) > 20:
        lines.append(f"... 共 {len(results)} 条，仅显示前 20 条")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════
# 干员档案文档生成工具
# ══════════════════════════════════════════════════════════════════

def _set_cell_shading(cell, color):
    """设置单元格背景色"""
    from docx.oxml.ns import qn
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)


def _extract_operator_info(soup):
    """从prts wiki页面提取干员信息"""
    info = {
        'name': '',
        'gender': '',
        'combat_experience': '',
        'birthplace': '',
        'birthday': '',
        'race': '',
        'height': '',
        'infection_status': '',
        'profession': '',
        'sub_profession': '',
        'redeploy_time': '',
        'deploy_cost': '',
        'block_count': '',
        'attack_interval': '',
        'faction': '',
    }

    # 获取页面文本
    text = soup.get_text(separator='\n')

    # 提取基本信息
    patterns = {
        'name': r'【代号】(.+)',
        'gender': r'【性别】(.+)',
        'combat_experience': r'【战斗经验】(.+)',
        'birthplace': r'【出身地】(.+)',
        'birthday': r'【生日】(.+)',
        'race': r'【种族】(.+)',
        'height': r'【身高】(.+)',
        'infection_status': r'【矿石病感染情况】(.+)',
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if match:
            info[key] = match.group(1).strip()

    # 提取属性信息
    attr_patterns = {
        'redeploy_time': r'再部署时间\s*(\d+s)',
        'deploy_cost': r'初始部署费用\s*(\d+→\d+)',
        'block_count': r'阻挡数\s*(\d+)',
        'attack_interval': r'攻击间隔\s*([\d.]+s)',
        'faction': r'所属势力\s*(.+)',
    }

    for key, pattern in attr_patterns.items():
        match = re.search(pattern, text)
        if match:
            info[key] = match.group(1).strip()

    return info


@mcp.tool()
def generate_operator_profile(operator_name: str, output_path: Optional[str] = None) -> str:
    """
    生成干员档案DOCX文档。

    参数:
        operator_name: 干员名称（如 "凯尔希·思衡托"、"Mon3tr"、"维什戴尔"）
        output_path: 保存路径（含文件名），为 None 时保存到当前目录

    返回:
        生成的DOCX文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        return "错误: 缺少 python-docx 库，请运行 'pip install python-docx' 安装"

    # 获取prts wiki页面
    url = f"https://prts.wiki/w/{operator_name}"
    try:
        res = requests.get(url, headers={"User-Agent": _MOBILE_UA}, timeout=15)
        res.encoding = 'utf-8'
        soup = BeautifulSoup(res.text, 'html.parser')
    except Exception as e:
        return f"错误: 无法获取页面 {e}"

    # 检查页面是否存在
    if '页面不存在' in res.text or '没有找到' in res.text:
        return f"错误: 未找到干员 '{operator_name}'"

    # 提取干员信息
    info = _extract_operator_info(soup)

    # 获取页面文本内容
    content = soup.find('div', {'id': 'mw-content-text'})
    if not content:
        return "错误: 无法获取页面内容"

    text = content.get_text(separator='\n')

    # 创建文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)

    # 标题
    title = doc.add_heading(f'{operator_name} 干员档案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    doc.add_heading('基本信息', level=1)

    # 基本信息表格
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ('代号', info['name'] or operator_name),
        ('性别', info['gender'] or '未知'),
        ('战斗经验', info['combat_experience'] or '未知'),
        ('出身地', info['birthplace'] or '未知'),
        ('生日', info['birthday'] or '未知'),
        ('种族', info['race'] or '未知'),
        ('身高', info['height'] or '未知'),
        ('矿石病感染情况', info['infection_status'] or '未知'),
    ]

    for i, (key, value) in enumerate(info_data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        _set_cell_shading(row.cells[0], 'E8E8E8')

    doc.add_paragraph()

    # 干员属性
    doc.add_heading('干员属性', level=1)

    # 从页面提取职业信息
    profession_match = re.search(r'职业\s*(\w+)', text)
    sub_profession_match = re.search(r'分支\s*(\w+)', text)

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    attr_data = [
        ('职业', profession_match.group(1) if profession_match else '未知'),
        ('分支', sub_profession_match.group(1) if sub_profession_match else '未知'),
        ('再部署时间', info['redeploy_time'] or '未知'),
        ('初始部署费用', info['deploy_cost'] or '未知'),
        ('阻挡数', info['block_count'] or '未知'),
        ('攻击间隔', info['attack_interval'] or '未知'),
        ('所属势力', info['faction'] or '未知'),
    ]

    for i, (key, value) in enumerate(attr_data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        _set_cell_shading(row.cells[0], 'E8E8E8')

    doc.add_paragraph()

    # 提取并添加天赋信息
    doc.add_heading('天赋', level=1)

    # 查找天赋部分
    talent_start = text.find('天赋')
    if talent_start >= 0:
        talent_text = text[talent_start:talent_start+2000]
        # 简化处理：提取天赋名称和描述
        talent_lines = talent_text.split('\n')
        current_talent = None
        for line in talent_lines[:30]:  # 限制行数
            line = line.strip()
            if not line:
                continue
            if '第一天赋' in line or '第二天赋' in line:
                current_talent = line
                doc.add_heading(line, level=2)
            elif current_talent and len(line) > 10:
                doc.add_paragraph(line)

    doc.add_paragraph()

    # 提取并添加技能信息
    doc.add_heading('技能', level=1)

    # 查找技能部分
    skill_start = text.find('技能')
    if skill_start >= 0:
        skill_text = text[skill_start:skill_start+3000]
        skill_lines = skill_text.split('\n')
        current_skill = None
        skill_count = 0
        for line in skill_lines[:50]:  # 限制行数
            line = line.strip()
            if not line:
                continue
            if '技能1' in line or '技能2' in line or '技能3' in line:
                current_skill = line
                skill_count += 1
                if skill_count <= 3:
                    doc.add_heading(line, level=2)
            elif current_skill and ('描述' in line or '攻击力' in line or '治疗' in line):
                if len(line) > 10:
                    doc.add_paragraph(line)

    doc.add_paragraph()

    # 保存文档
    if output_path:
        save_path = output_path
    else:
        save_path = f"{operator_name} 干员档案.docx"
    os.makedirs(os.path.dirname(os.path.abspath(save_path)), exist_ok=True)
    doc.save(save_path)

    # 返回完整路径
    full_path = os.path.abspath(save_path)
    return f"文档已成功生成：{full_path}"


# ══════════════════════════════════════════════════════════════════
# 文档编辑工具
# ══════════════════════════════════════════════════════════════════

@mcp.tool()
def edit_document(
    file_path: str,
    operation: str,
    old_text: str = "",
    new_text: str = "",
    text: str = "",
    heading: str = "",
    content: str = "",
    level: int = 1,
    position: int = -1,
    row: int = 0,
    col: int = 0,
    table_index: int = 0,
    font_name: str = "",
    font_size: int = 0,
) -> str:
    """
    创建或编辑 DOCX 文档。文件不存在时自动创建新文档。

    参数:
        file_path:    DOCX 文件路径
        operation:    操作类型：
                      - create:        创建新文档（如已存在则覆盖），可通过 content 传入初始内容
                      - read:          读取文档结构（段落和表格内容概览）
                      - replace_text:  查找替换文本
                      - add_paragraph: 在指定位置插入段落
                      - add_heading:   在指定位置插入标题
                      - add_section:   在指定位置插入标题+段落组合
                      - add_table:     插入表格（content 为 JSON 二维数组）
                      - edit_table:    修改表格单元格内容
                      - set_font:      设置文档默认字体
        old_text:     replace_text 用：要查找的文本
        new_text:     replace_text 用：替换后的文本 / edit_table 用：单元格新值
        text:         add_paragraph / add_heading 用：文本内容
        heading:      add_section 用：标题文本
        content:      add_section 用：段落内容（可用 \\n 分隔多段）
                      create 用：初始正文内容（可用 \\n 分段）
                      add_table 用：JSON 二维数组，如 '[["姓名","年龄"],["Mon3tr","?"]]'
        level:        add_heading 用：标题级别（1-4，默认 1）
        position:     插入位置索引（默认 -1 即文档末尾，0 为文档开头）
        row:          edit_table 用：行索引（从 0 开始）
        col:          edit_table 用：列索引（从 0 开始）
        table_index:  edit_table 用：第几个表格（默认第 1 个，从 0 开始）
        font_name:    字体名称（如 "宋体"、"Times New Roman"），create/set_font 时使用，默认微软雅黑
        font_size:    字号（整数，单位磅），create/set_font 时使用，默认 10.5
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except ImportError:
        return "错误: 缺少 python-docx 库，请运行 'pip install python-docx' 安装"

    abs_path = os.path.abspath(file_path)
    os.makedirs(os.path.dirname(abs_path) or ".", exist_ok=True)

    _font = font_name or "微软雅黑"
    _size = Pt(font_size) if font_size > 0 else Pt(10.5)

    def _apply_font(doc):
        style = doc.styles['Normal']
        style.font.name = _font
        style.element.rPr.rFonts.set(qn('w:eastAsia'), _font)
        style.font.size = _size

    # ── create: 创建新文档 ──
    if operation == "create":
        doc = Document()
        _apply_font(doc)
        if text:
            doc.add_heading(text, level=0)
        if content:
            for para_text in content.split("\n"):
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)
        doc.save(abs_path)
        return f"文档已创建：{abs_path}"

    # 非 create 操作：打开已有文档或创建空文档
    if os.path.isfile(abs_path):
        try:
            doc = Document(abs_path)
        except Exception as e:
            return f"错误: 无法打开文档 {e}"
    else:
        doc = Document()
        _apply_font(doc)

    # ── set_font: 修改文档默认字体 ──
    if operation == "set_font":
        _apply_font(doc)
        # 同时更新已有段落的字体
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = _font
                run.element.rPr.rFonts.set(qn('w:eastAsia'), _font)
                if font_size > 0:
                    run.font.size = _size
        doc.save(abs_path)
        return f"已设置字体为 {_font}，字号 {font_size or 10.5}pt"

    # ── read: 读取文档结构 ──
    if operation == "read":
        lines = [f"文档: {abs_path}", ""]
        for i, para in enumerate(doc.paragraphs):
            style = para.style.name if para.style else ""
            prefix = f"[{i}] "
            if "Heading" in style:
                level_num = style.replace("Heading ", "")
                prefix += f"(H{level_num}) "
            text_content = para.text.strip()
            if text_content:
                lines.append(f"{prefix}{text_content[:120]}")
            else:
                lines.append(f"{prefix}(空行)")
        for ti, table in enumerate(doc.tables):
            lines.append(f"\n── 表格 {ti} ({len(table.rows)}行 x {len(table.columns)}列) ──")
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip()[:30] for cell in row.cells]
                lines.append(f"  [{ri}] | {' | '.join(cells)} |")
        return "\n".join(lines)

    # ── replace_text: 查找替换 ──
    if operation == "replace_text":
        if not old_text:
            return "错误: 请提供 old_text 参数"
        count = 0
        for para in doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        count += 1
        for table in doc.tables:
            for row_obj in table.rows:
                for cell in row_obj.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                                    count += 1
        doc.save(abs_path)
        return f"替换完成，共替换 {count} 处" if count > 0 else "未找到匹配文本"

    # ── add_paragraph: 插入段落 ──
    if operation == "add_paragraph":
        if not text:
            return "错误: 请提供 text 参数"
        body = doc.element.body
        new_para = doc.add_paragraph(text)
        if position >= 0 and position < len(body):
            body.insert(position, new_para._element)
        doc.save(abs_path)
        return f"已在位置 {position} 插入段落"

    # ── add_heading: 插入标题 ──
    if operation == "add_heading":
        if not text:
            return "错误: 请提供 text 参数"
        level = max(1, min(4, level))
        body = doc.element.body
        new_heading = doc.add_heading(text, level=level)
        if position >= 0 and position < len(body):
            body.insert(position, new_heading._element)
        doc.save(abs_path)
        return f"已在位置 {position} 插入 H{level} 标题"

    # ── add_section: 插入标题+段落组合 ──
    if operation == "add_section":
        if not heading and not content:
            return "错误: 请提供 heading 或 content 参数"
        body = doc.element.body
        insert_pos = position if position >= 0 else len(body)
        if heading:
            h = doc.add_heading(heading, level=level)
            body.insert(insert_pos, h._element)
            insert_pos += 1
        if content:
            for para_text in content.split("\n"):
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    body.insert(insert_pos, p._element)
                    insert_pos += 1
        doc.save(abs_path)
        return f"已在位置 {position} 插入章节"

    # ── add_table: 插入表格 ──
    if operation == "add_table":
        if not content:
            return "错误: 请提供 content 参数（JSON 二维数组）"
        try:
            table_data = json.loads(content)
        except json.JSONDecodeError as e:
            return f"错误: content 不是有效 JSON: {e}"
        if not table_data or not isinstance(table_data, list):
            return "错误: content 应为二维数组"
        rows_count = len(table_data)
        cols_count = max(len(r) for r in table_data)
        table = doc.add_table(rows=rows_count, cols=cols_count)
        table.style = 'Table Grid'
        for ri, row_data in enumerate(table_data):
            for ci, val in enumerate(row_data):
                table.rows[ri].cells[ci].text = str(val)
        doc.save(abs_path)
        return f"已插入 {rows_count}x{cols_count} 表格"

    # ── edit_table: 修改表格单元格 ──
    if operation == "edit_table":
        tables = doc.tables
        if table_index < 0 or table_index >= len(tables):
            return f"错误: 表格索引 {table_index} 超出范围（共 {len(tables)} 个表格）"
        table = tables[table_index]
        if row < 0 or row >= len(table.rows):
            return f"错误: 行索引 {row} 超出范围（共 {len(table.rows)} 行）"
        if col < 0 or col >= len(table.columns):
            return f"错误: 列索引 {col} 超出范围（共 {len(table.columns)} 列）"
        cell = table.rows[row].cells[col]
        old_val = cell.text.strip()
        cell.text = new_text
        doc.save(abs_path)
        return f"表格{table_index}[{row},{col}]: '{old_val}' -> '{new_text}'"

    return f"错误: 未知操作 '{operation}'，支持: create, read, replace_text, add_paragraph, add_heading, add_section, add_table, edit_table, set_font"


# ══════════════════════════════════════════════════════════════════
# Excel 表格工具
# ══════════════════════════════════════════════════════════════════

@mcp.tool()
def edit_excel(
    file_path: str,
    operation: str,
    sheet: str = "",
    cell: str = "",
    value: str = "",
    data: str = "",
    formula_str: str = "",
    bold: bool = False,
    font_color: str = "",
    bg_color: str = "",
    font_size: int = 0,
    align: str = "",
    border: bool = False,
    col_width: str = "",
    row_height: str = "",
    params: str = "",
) -> str:
    """
    创建和编辑 Excel (.xlsx) 表格。

    参数:
        file_path:    Excel 文件路径
        operation:    操作类型：
                      - create:            创建新文件（已有则覆盖）
                      - read:              读取工作表内容
                      - write:             写入数据（单格或批量）
                      - formula:           写入公式
                      - add_sheet:         添加新工作表
                      - format:            设置单元格格式
                      - merge:             合并单元格
                      - col_width:         设置列宽（如 "A:20,B:15"）
                      - row_height:        设置行高（如 "1:30,2:20"）
                      - chart:             添加图表
                      - conditional_format: 条件格式（色阶/数据条）
                      - data_validation:   数据验证（下拉列表）
                      - freeze_panes:      冻结窗格
                      - auto_filter:       自动筛选
                      - image:             插入图片
        sheet:        工作表名称（默认活动表）
        cell:         单元格引用，如 "A1" 或范围 "A1:C3"
        value:        单格写入的值
        data:         批量写入数据，JSON 二维数组，如 "[["姓名","分数"],["张三",90]]"
        formula_str:  公式（不含=号），如 "SUM(A1:A10)"
        bold:         是否加粗
        font_color:   字体颜色，如 "FF0000"
        bg_color:     背景色，如 "FFFF00"
        font_size:    字号
        align:        对齐方式：left / center / right
        border:       是否加边框
        col_width:    操作类型为 col_width 时：列宽设置，如 "A:20,B:15"
        row_height:   操作类型为 row_height 时：行高设置，如 "1:30,2:20"
        params:       高级操作的 JSON 参数，各操作可用字段：
                      chart:             {"type":"bar","title":"标题","data":"B1:D4","cats":"A2:A4","pos":"F2"}
                      conditional_format: {"range":"B2:D4","type":"color_scale","colors":["63BE7B","FFEB84","F8696B"]}
                      data_validation:   {"range":"B2:B100","list":"选项1,选项2,选项3"}
                      freeze_panes:      {"cell":"B2"}
                      auto_filter:       {"range":"A1:D10"}
                      image:             {"path":"logo.png","anchor":"A1","width":200,"height":100}
    """
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import range_boundaries
    except ImportError:
        return "错误: 缺少 openpyxl 库，请运行 'pip install openpyxl' 安装"

    abs_path = os.path.abspath(file_path)

    # ── helper ──
    def _get_sheet(wb):
        if sheet and sheet in wb.sheetnames:
            return wb[sheet]
        return wb.active

    def _apply_format(ws, cell_ref):
        """对单格或范围应用格式"""
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import range_boundaries

        if ":" in cell_ref:
            min_col, min_row, max_col, max_row = range_boundaries(cell_ref)
            targets = []
            for r in range(min_row, max_row + 1):
                for c in range(min_col, max_col + 1):
                    targets.append(ws.cell(row=r, column=c))
        else:
            targets = [ws[cell_ref]]

        for c in targets:
            if bold:
                c.font = Font(bold=True, size=font_size or None, color=font_color or None)
            elif font_size or font_color:
                c.font = Font(size=font_size or None, color=font_color or None)
            if bg_color:
                c.fill = PatternFill(start_color=bg_color, end_color=bg_color, fill_type="solid")
            if align:
                c.alignment = Alignment(horizontal=align)
            if border:
                thin = Side(style="thin")
                c.border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── create: 创建新文件 ──
    if operation == "create":
        wb = Workbook()
        ws = wb.active
        ws.title = sheet or "Sheet1"
        if data:
            try:
                rows = json.loads(data)
                for r, row_data in enumerate(rows, 1):
                    for c, val in enumerate(row_data, 1):
                        ws.cell(row=r, column=c, value=val)
            except json.JSONDecodeError:
                return "错误: data 不是有效的 JSON 二维数组"
        os.makedirs(os.path.dirname(abs_path) or ".", exist_ok=True)
        wb.save(abs_path)
        sheets_info = f"工作表: {ws.title}"
        if data:
            rows = json.loads(data)
            sheets_info += f"，{len(rows)} 行数据"
        return f"已创建 {abs_path}\n{sheets_info}"

    # ── 以下操作需要打开已有文件 ──
    if not os.path.isfile(abs_path):
        return f"错误: 文件不存在 {abs_path}"

    # ── read: 读取内容 ──
    if operation == "read":
        wb = load_workbook(abs_path, data_only=True)
        ws = _get_sheet(wb)
        lines = [f"文件: {abs_path}", f"工作表: {ws.title}", f"尺寸: {ws.max_row} 行 x {ws.max_column} 列", ""]

        if cell and ":" in cell:
            target = ws[cell]
            from openpyxl.utils import range_boundaries
            min_col, min_row, max_col, max_row = range_boundaries(cell)
            for r in range(min_row, max_row + 1):
                vals = []
                for c in range(min_col, max_col + 1):
                    v = ws.cell(row=r, column=c).value
                    vals.append(str(v) if v is not None else "")
                lines.append(f"  [{r}] | {' | '.join(vals)} |")
        else:
            for r in range(1, min(ws.max_row + 1, 101)):
                vals = []
                for c in range(1, min(ws.max_column + 1, 27)):
                    v = ws.cell(row=r, column=c).value
                    vals.append(str(v) if v is not None else "")
                line = f"  [{r}] | {' | '.join(vals)} |"
                lines.append(line)
            if ws.max_row > 100:
                lines.append(f"  ... 共 {ws.max_row} 行，仅显示前 100 行")
        return "\n".join(lines)

    # ── write: 写入数据 ──
    if operation == "write":
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        count = 0
        if data:
            try:
                rows_data = json.loads(data)
            except json.JSONDecodeError:
                return "错误: data 不是有效的 JSON 二维数组"
            if cell:
                from openpyxl.utils.cell import coordinate_to_tuple
                start_row, start_col = coordinate_to_tuple(cell)
            else:
                start_row, start_col = 1, 1
            for r, row_data in enumerate(rows_data):
                for c, val in enumerate(row_data):
                    ws.cell(row=start_row + r, column=start_col + c, value=val)
                    count += 1
        elif cell and value != "":
            if ":" in cell:
                from openpyxl.utils import range_boundaries
                min_col, min_row, max_col, max_row = range_boundaries(cell)
                for r in range(min_row, max_row + 1):
                    for c in range(min_col, max_col + 1):
                        ws.cell(row=r, column=c, value=value)
                        count += 1
            else:
                ws[cell] = value
                count = 1
        else:
            return "错误: 请提供 cell+value 或 data 参数"
        wb.save(abs_path)
        return f"已写入 {count} 个单元格"

    # ── formula: 写入公式 ──
    if operation == "formula":
        if not cell or not formula_str:
            return "错误: 请提供 cell 和 formula_str 参数"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        ws[cell] = f"={formula_str}"
        wb.save(abs_path)
        return f"已写入公式: {cell} = {formula_str}"

    # ── add_sheet: 添加工作表 ──
    if operation == "add_sheet":
        if not sheet:
            return "错误: 请提供 sheet 参数（新工作表名称）"
        wb = load_workbook(abs_path)
        if sheet in wb.sheetnames:
            return f"错误: 工作表 '{sheet}' 已存在"
        wb.create_sheet(title=sheet)
        wb.save(abs_path)
        return f"已添加工作表 '{sheet}'，当前工作表: {', '.join(wb.sheetnames)}"

    # ── format: 设置格式 ──
    if operation == "format":
        if not cell:
            return "错误: 请提供 cell 参数"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        _apply_format(ws, cell)
        wb.save(abs_path)
        return f"已设置 {cell} 的格式"

    # ── merge: 合并单元格 ──
    if operation == "merge":
        if not cell or ":" not in cell:
            return "错误: 请提供 cell 参数（范围如 A1:C3）"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        ws.merge_cells(cell)
        wb.save(abs_path)
        return f"已合并 {cell}"

    # ── col_width: 设置列宽 ──
    if operation == "col_width":
        if not col_width:
            return "错误: 请提供 col_width 参数（如 'A:20,B:15'）"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        for item in col_width.split(","):
            item = item.strip()
            if ":" in item:
                col_letter, width = item.split(":", 1)
                ws.column_dimensions[col_letter.strip()].width = float(width.strip())
        wb.save(abs_path)
        return f"已设置列宽: {col_width}"

    # ── row_height: 设置行高 ──
    if operation == "row_height":
        if not row_height:
            return "错误: 请提供 row_height 参数（如 '1:30,2:20'）"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        for item in row_height.split(","):
            item = item.strip()
            if ":" in item:
                row_num, height = item.split(":", 1)
                ws.row_dimensions[int(row_num)].height = float(height.strip())
        wb.save(abs_path)
        return f"已设置行高: {row_height}"

    # ── 解析 params JSON ──
    p = {}
    if params:
        try:
            p = json.loads(params)
        except json.JSONDecodeError:
            return "错误: params 不是有效的 JSON"

    # ── chart: 添加图表 ──
    if operation == "chart":
        from openpyxl.chart import BarChart, LineChart, PieChart, AreaChart, ScatterChart, Reference
        chart_map = {
            "bar": BarChart, "line": LineChart, "pie": PieChart,
            "area": AreaChart, "scatter": ScatterChart,
        }
        ctype = p.get("type", "bar")
        if ctype not in chart_map:
            return f"错误: 不支持的图表类型 '{ctype}'，支持: {', '.join(chart_map)}"
        if not cell:
            return "错误: 请提供 cell 参数指定数据范围（如 A1:D4）"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        chart = chart_map[ctype]()
        chart.title = p.get("title", "")
        chart.style = int(p.get("style", 10))
        # 数据系列
        from openpyxl.utils import range_boundaries as rb
        d_min_col, d_min_row, d_max_col, d_max_row = rb(cell)
        cats_ref = p.get("cats", "")
        if cats_ref:
            cats = Reference(ws, *rb(cats_ref)[0:2], *rb(cats_ref)[2:4])
        else:
            cats = Reference(ws, min_col=d_min_col, min_row=d_min_row + 1, max_row=d_max_row)
        for c in range(d_min_col + (1 if not cats_ref else 0), d_max_col + 1):
            vals = Reference(ws, min_col=c, min_row=d_min_row, max_row=d_max_row)
            chart.add_data(vals, titles_from_data=True)
        chart.set_categories(cats)
        # 尺寸
        chart.width = float(p.get("width", 15))
        chart.height = float(p.get("height", 10))
        # 位置
        pos = p.get("pos", "F2")
        ws.add_chart(chart, pos)
        wb.save(abs_path)
        return f"已添加 {ctype} 图表，数据范围 {cell}，位置 {pos}"

    # ── conditional_format: 条件格式 ──
    if operation == "conditional_format":
        from openpyxl.formatting.rule import ColorScaleRule, DataBarRule
        from openpyxl.utils import range_boundaries as rb
        if not cell:
            return "错误: 请提供 cell 参数指定范围（如 B2:D10）"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        cf_type = p.get("type", "color_scale")
        colors = p.get("colors", ["63BE7B", "FFEB84", "F8696B"])
        if cf_type == "color_scale":
            rule = ColorScaleRule(
                start_type="min", start_color=colors[0],
                mid_type="percentile", mid_value=50, mid_color=colors[1] if len(colors) > 1 else "FFFFFF",
                end_type="max", end_color=colors[2] if len(colors) > 2 else "F8696B",
            )
            ws.conditional_formatting.add(cell, rule)
        elif cf_type == "data_bar":
            from openpyxl.formatting.rule import DataBarRule
            rule = DataBarRule(start_type="min", end_type="max", color=colors[0])
            ws.conditional_formatting.add(cell, rule)
        else:
            return f"错误: 不支持的条件格式类型 '{cf_type}'，支持: color_scale, data_bar"
        wb.save(abs_path)
        return f"已对 {cell} 添加 {cf_type} 条件格式"

    # ── data_validation: 数据验证 ──
    if operation == "data_validation":
        from openpyxl.worksheet.datavalidation import DataValidation
        if not cell:
            return "错误: 请提供 cell 参数指定范围（如 B2:B100）"
        list_str = p.get("list", "")
        if not list_str:
            return "错误: params 中请提供 list 字段（逗号分隔的选项）"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        dv = DataValidation(type="list", formula1=f'"{list_str}"', allow_blank=True)
        dv.error = "请从列表中选择"
        dv.errorTitle = "无效输入"
        ws.add_data_validation(dv)
        dv.add(cell)
        wb.save(abs_path)
        return f"已对 {cell} 添加下拉列表: {list_str}"

    # ── freeze_panes: 冻结窗格 ──
    if operation == "freeze_panes":
        freeze_cell = p.get("cell", cell or "A2")
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        ws.freeze_panes = freeze_cell
        wb.save(abs_path)
        return f"已冻结窗格: {freeze_cell}"

    # ── auto_filter: 自动筛选 ──
    if operation == "auto_filter":
        filter_range = cell or p.get("range", "")
        if not filter_range:
            return "错误: 请提供 cell 参数指定筛选范围（如 A1:D10）"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        ws.auto_filter.ref = filter_range
        wb.save(abs_path)
        return f"已对 {filter_range} 启用自动筛选"

    # ── image: 插入图片 ──
    if operation == "image":
        from openpyxl.drawing.image import Image as XlImage
        img_path = p.get("path", "")
        if not img_path:
            return "错误: params 中请提供 path 字段（图片路径）"
        img_abs = os.path.abspath(img_path)
        if not os.path.isfile(img_abs):
            return f"错误: 图片不存在 {img_abs}"
        wb = load_workbook(abs_path)
        ws = _get_sheet(wb)
        img = XlImage(img_abs)
        if "width" in p:
            img.width = int(p["width"])
        if "height" in p:
            img.height = int(p["height"])
        anchor = p.get("anchor", cell or "A1")
        ws.add_image(img, anchor)
        wb.save(abs_path)
        return f"已插入图片 {img_path}，锚点 {anchor}"

    return (
        f"错误: 未知操作 '{operation}'，支持: "
        "create, read, write, formula, add_sheet, format, merge, col_width, row_height, "
        "chart, conditional_format, data_validation, freeze_panes, auto_filter, image"
    )


# ══════════════════════════════════════════════════════════════════
# PDF 读写工具
# ══════════════════════════════════════════════════════════════════

@mcp.tool()
def edit_pdf(
    file_path: str,
    operation: str,
    output_path: str = "",
    pages: str = "",
    content: str = "",
    title: str = "",
    author: str = "",
    font_size: int = 12,
    page_size: str = "A4",
    margin: int = 72,
    params: str = "",
) -> str:
    """
    读取和生成 PDF 文件。

    参数:
        file_path:    PDF 文件路径（read/info/merge/extract 用）
        operation:    操作类型：
                      - read:    读取 PDF 文本内容
                      - info:    获取 PDF 元信息（页数、标题、作者等）
                      - create:  创建新 PDF（从文本生成）
                      - merge:   合并多个 PDF
                      - extract: 提取指定页面到新文件
                      - to_word: PDF 转 Word (.docx)
        output_path:  输出文件路径（create/merge/extract 必填）
        pages:        页码范围，如 "1,3,5-8"（read/extract 用，默认全部）
        content:      文本内容（create 用，用 \\n 分段）
        title:        文档标题（create 用）
        author:       文档作者（create 用）
        font_size:    字号（create 用，默认 12）
        page_size:    页面大小（create 用：A4/Letter/A5，默认 A4）
        margin:       页边距点数（create 用，默认 72，即 1 英寸）
        params:       JSON 参数，高级选项：
                      merge:   {"files":["a.pdf","b.pdf","c.pdf"]}
                      create:  {"line_spacing":1.5,"header":"页眉","footer":"页脚"}
    """
    abs_path = os.path.abspath(file_path)

    # ── helper: 解析页码范围 ──
    def _parse_pages(pages_str: str, max_page: int) -> list:
        if not pages_str:
            return list(range(max_page))
        result = []
        for part in pages_str.split(","):
            part = part.strip()
            if "-" in part:
                start, end = part.split("-", 1)
                result.extend(range(int(start) - 1, int(end)))
            else:
                result.append(int(part) - 1)
        return [p for p in result if 0 <= p < max_page]

    # ── read: 读取文本 ──
    if operation == "read":
        try:
            from pypdf import PdfReader
        except ImportError:
            return "错误: 缺少 pypdf 库，请运行 'pip install pypdf' 安装"
        if not os.path.isfile(abs_path):
            return f"错误: 文件不存在 {abs_path}"
        reader = PdfReader(abs_path)
        page_list = _parse_pages(pages, len(reader.pages))
        lines = [f"文件: {abs_path}", f"总页数: {len(reader.pages)}", ""]
        for i in page_list:
            text = reader.pages[i].extract_text() or ""
            lines.append(f"── 第 {i + 1} 页 ──")
            lines.append(text.strip() if text.strip() else "(无可提取文本)")
            lines.append("")
        return "\n".join(lines)

    # ── info: 获取元信息 ──
    if operation == "info":
        try:
            from pypdf import PdfReader
        except ImportError:
            return "错误: 缺少 pypdf 库，请运行 'pip install pypdf' 安装"
        if not os.path.isfile(abs_path):
            return f"错误: 文件不存在 {abs_path}"
        reader = PdfReader(abs_path)
        meta = reader.metadata
        lines = [
            f"文件: {abs_path}",
            f"页数: {len(reader.pages)}",
            f"标题: {meta.title if meta and meta.title else '无'}",
            f"作者: {meta.author if meta and meta.author else '无'}",
            f"主题: {meta.subject if meta and meta.subject else '无'}",
            f"创建者: {meta.creator if meta and meta.creator else '无'}",
        ]
        if reader.pages:
            page = reader.pages[0]
            box = page.mediabox
            lines.append(f"页面尺寸: {float(box.width):.0f} x {float(box.height):.0f} 点")
        lines.append(f"加密: {'是' if reader.is_encrypted else '否'}")
        return "\n".join(lines)

    # ── create: 创建 PDF ──
    if operation == "create":
        try:
            from reportlab.lib.pagesizes import A4, LETTER, A5
            from reportlab.lib.units import mm
            from reportlab.pdfgen import canvas
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            return "错误: 缺少 reportlab 库，请运行 'pip install reportlab' 安装"
        if not content:
            return "错误: 请提供 content 参数"
        if not output_path:
            return "错误: 请提供 output_path 参数"

        size_map = {"A4": A4, "Letter": LETTER, "A5": A5}
        pagesize = size_map.get(page_size, A4)

        # 解析 params
        p = {}
        if params:
            try:
                p = json.loads(params)
            except json.JSONDecodeError:
                return "错误: params 不是有效的 JSON"

        out_abs = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(out_abs) or ".", exist_ok=True)

        # 注册中文字体（尝试常见路径）
        chinese_font = None
        font_paths = [
            "C:/Windows/Fonts/msyh.ttc",      # 微软雅黑
            "C:/Windows/Fonts/simsun.ttc",      # 宋体
            "C:/Windows/Fonts/simhei.ttf",      # 黑体
            "/System/Library/Fonts/PingFang.ttc",  # macOS
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # Linux
        ]
        for fp in font_paths:
            if os.path.isfile(fp):
                try:
                    pdfmetrics.registerFont(TTFont("ChineseFont", fp))
                    chinese_font = "ChineseFont"
                    break
                except Exception:
                    continue

        font_name = chinese_font or "Helvetica"

        doc = SimpleDocTemplate(
            out_abs, pagesize=pagesize,
            leftMargin=margin, rightMargin=margin,
            topMargin=margin, bottomMargin=margin,
            title=title or "",
            author=author or "",
        )

        style = ParagraphStyle(
            "Body",
            fontName=font_name,
            fontSize=font_size,
            leading=font_size * float(p.get("line_spacing", 1.5)),
        )
        title_style = ParagraphStyle(
            "Title",
            fontName=font_name,
            fontSize=font_size + 6,
            leading=(font_size + 6) * 1.5,
            alignment=1,  # center
            spaceAfter=20,
        )

        story = []
        if title:
            story.append(Paragraph(title, title_style))
        for para in content.split("\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para, style))
                story.append(Spacer(1, font_size * 0.5))

        doc.build(story)
        return f"已创建 PDF: {out_abs}（{len(content.split(chr(10)))} 段）"

    # ── merge: 合并 PDF ──
    if operation == "merge":
        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError:
            return "错误: 缺少 pypdf 库，请运行 'pip install pypdf' 安装"
        if not output_path:
            return "错误: 请提供 output_path 参数"

        p = {}
        if params:
            try:
                p = json.loads(params)
            except json.JSONDecodeError:
                return "错误: params 不是有效的 JSON"

        files = p.get("files", [])
        if not files and file_path:
            files = [file_path]
        if not files:
            return "错误: 请在 params.files 中提供要合并的 PDF 文件列表"

        writer = PdfWriter()
        total_pages = 0
        for f in files:
            f_abs = os.path.abspath(f)
            if not os.path.isfile(f_abs):
                return f"错误: 文件不存在 {f_abs}"
            reader = PdfReader(f_abs)
            for page in reader.pages:
                writer.add_page(page)
                total_pages += 1

        out_abs = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(out_abs) or ".", exist_ok=True)
        with open(out_abs, "wb") as f:
            writer.write(f)
        return f"已合并 {len(files)} 个 PDF（共 {total_pages} 页）-> {out_abs}"

    # ── extract: 提取页面 ──
    if operation == "extract":
        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError:
            return "错误: 缺少 pypdf 库，请运行 'pip install pypdf' 安装"
        if not os.path.isfile(abs_path):
            return f"错误: 文件不存在 {abs_path}"
        if not output_path:
            return "错误: 请提供 output_path 参数"
        reader = PdfReader(abs_path)
        page_list = _parse_pages(pages, len(reader.pages))
        if not page_list:
            return "错误: 未指定有效页码"
        writer = PdfWriter()
        for i in page_list:
            writer.add_page(reader.pages[i])
        out_abs = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(out_abs) or ".", exist_ok=True)
        with open(out_abs, "wb") as f:
            writer.write(f)
        return f"已提取 {len(page_list)} 页 -> {out_abs}"

    # ── to_word: PDF 转 Word ──
    if operation == "to_word":
        try:
            from pdf2docx import Converter
        except ImportError:
            return "错误: 缺少 pdf2docx 库，请运行 'pip install pdf2docx' 安装"
        if not os.path.isfile(abs_path):
            return f"错误: 文件不存在 {abs_path}"
        out_abs = os.path.abspath(output_path) if output_path else abs_path.rsplit(".", 1)[0] + ".docx"
        os.makedirs(os.path.dirname(out_abs) or ".", exist_ok=True)
        # 解析页码范围
        p = {}
        if params:
            try:
                p = json.loads(params)
            except json.JSONDecodeError:
                return "错误: params 不是有效的 JSON"
        cv = Converter(abs_path)
        start_page = p.get("start", 0)
        end_page = p.get("end", None)
        if pages:
            page_list = _parse_pages(pages, len(cv.pages))
            if page_list:
                start_page = page_list[0]
                end_page = page_list[-1] + 1
        cv.convert(out_abs, start=start_page, end=end_page)
        cv.close()
        return f"已转换: {abs_path} -> {out_abs}"

    return (
        f"错误: 未知操作 '{operation}'，支持: read, info, create, merge, extract, to_word"
    )


# ══════════════════════════════════════════════════════════════════
# TXT 文本文件工具
# ══════════════════════════════════════════════════════════════════

@mcp.tool()
def edit_txt(
    file_path: str,
    operation: str,
    content: str = "",
    encoding: str = "utf-8",
    old_text: str = "",
    new_text: str = "",
    line: int = 0,
    count: int = 0,
) -> str:
    """
    创建和编辑纯文本 (.txt) 文件。

    参数:
        file_path:  TXT 文件路径
        operation:  操作类型：
                    - create:   创建新文件（已有则覆盖）
                    - read:     读取文件内容
                    - write:    写入内容（覆盖）
                    - append:   追加内容到末尾
                    - replace:  查找替换文本
                    - insert:   在指定行插入内容
                    - delete:   删除指定行
        content:    写入/追加/插入的内容
        encoding:   文件编码（默认 utf-8）
        old_text:   replace 用：要查找的文本
        new_text:   replace 用：替换后的文本
        line:       insert/delete 用：行号（从 1 开始）
        count:      delete 用：要删除的行数（默认 1）
    """
    abs_path = os.path.abspath(file_path)

    # ── create: 创建新文件 ──
    if operation == "create":
        os.makedirs(os.path.dirname(abs_path) or ".", exist_ok=True)
        with open(abs_path, "w", encoding=encoding) as f:
            f.write(content)
        return f"已创建 {abs_path}（{len(content)} 字符）"

    # ── read: 读取内容 ──
    if operation == "read":
        if not os.path.isfile(abs_path):
            return f"错误: 文件不存在 {abs_path}"
        with open(abs_path, "r", encoding=encoding) as f:
            text = f.read()
        lines_count = text.count("\n") + 1
        return f"文件: {abs_path}\n行数: {lines_count}\n字符数: {len(text)}\n\n{text}"

    # ── 以下操作需要文件存在 ──
    if not os.path.isfile(abs_path):
        return f"错误: 文件不存在 {abs_path}"

    # ── write: 写入内容（覆盖）──
    if operation == "write":
        with open(abs_path, "w", encoding=encoding) as f:
            f.write(content)
        return f"已写入 {abs_path}（{len(content)} 字符）"

    # ── append: 追加内容 ──
    if operation == "append":
        with open(abs_path, "a", encoding=encoding) as f:
            f.write(content)
        return f"已追加 {len(content)} 字符到 {abs_path}"

    # ── replace: 查找替换 ──
    if operation == "replace":
        if not old_text:
            return "错误: 请提供 old_text 参数"
        with open(abs_path, "r", encoding=encoding) as f:
            text = f.read()
        count = text.count(old_text)
        if count == 0:
            return "未找到匹配文本"
        text = text.replace(old_text, new_text)
        with open(abs_path, "w", encoding=encoding) as f:
            f.write(text)
        return f"替换完成，共替换 {count} 处"

    # ── insert: 在指定行插入 ──
    if operation == "insert":
        if line < 1:
            return "错误: 行号必须 >= 1"
        with open(abs_path, "r", encoding=encoding) as f:
            lines = f.readlines()
        insert_idx = min(line - 1, len(lines))
        new_lines = content.split("\n")
        for i, ln in enumerate(new_lines):
            lines.insert(insert_idx + i, ln + "\n")
        with open(abs_path, "w", encoding=encoding) as f:
            f.writelines(lines)
        return f"已在第 {line} 行插入 {len(new_lines)} 行"

    # ── delete: 删除指定行 ──
    if operation == "delete":
        if line < 1:
            return "错误: 行号必须 >= 1"
        with open(abs_path, "r", encoding=encoding) as f:
            lines = f.readlines()
        if line > len(lines):
            return f"错误: 行号 {line} 超出范围（共 {len(lines)} 行）"
        del_count = min(count or 1, len(lines) - line + 1)
        del lines[line - 1:line - 1 + del_count]
        with open(abs_path, "w", encoding=encoding) as f:
            f.writelines(lines)
        return f"已删除第 {line} 行起 {del_count} 行"

    return f"错误: 未知操作 '{operation}'，支持: create, read, write, append, replace, insert, delete"


# ══════════════════════════════════════════════════════════════════
# JSON 文件工具
# ══════════════════════════════════════════════════════════════════

@mcp.tool()
def edit_json(
    file_path: str,
    operation: str,
    data: str = "",
    key: str = "",
    value: str = "",
    indent: int = 2,
    merge_data: str = "",
) -> str:
    """
    创建和编辑 JSON (.json) 文件。

    参数:
        file_path:    JSON 文件路径
        operation:    操作类型：
                      - create:   创建新 JSON 文件（已有则覆盖）
                      - read:     读取 JSON 内容
                      - write:    写入/更新键值
                      - delete:   删除指定键
                      - merge:    合并另一个 JSON 对象
                      - array:    向数组追加元素
                      - validate: 验证 JSON 格式
        data:         create 用：JSON 字符串（如 '{"key": "value"}'）
                      array 用：要追加的 JSON 值
        key:          write/delete 用：键名（支持点号路径如 "a.b.c"）
        value:        write 用：值（JSON 字符串或纯文本）
        indent:       缩进空格数（默认 2）
        merge_data:   merge 用：要合并的 JSON 字符串
    """
    abs_path = os.path.abspath(file_path)

    def _set_nested(obj, path, val):
        """设置嵌套键值"""
        keys = path.split(".")
        for k in keys[:-1]:
            if k not in obj or not isinstance(obj[k], dict):
                obj[k] = {}
            obj = obj[k]
        obj[keys[-1]] = val

    def _get_nested(obj, path):
        """获取嵌套键值"""
        keys = path.split(".")
        for k in keys:
            if isinstance(obj, dict) and k in obj:
                obj = obj[k]
            else:
                return None
        return obj

    def _delete_nested(obj, path):
        """删除嵌套键"""
        keys = path.split(".")
        for k in keys[:-1]:
            if isinstance(obj, dict) and k in obj:
                obj = obj[k]
            else:
                return False
        if isinstance(obj, dict) and keys[-1] in obj:
            del obj[keys[-1]]
            return True
        return False

    # ── create: 创建新文件 ──
    if operation == "create":
        try:
            json_data = json.loads(data) if data else {}
        except json.JSONDecodeError as e:
            return f"错误: JSON 格式无效 - {e}"
        os.makedirs(os.path.dirname(abs_path) or ".", exist_ok=True)
        with open(abs_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=indent)
        return f"已创建 {abs_path}"

    # ── read: 读取内容 ──
    if operation == "read":
        if not os.path.isfile(abs_path):
            return f"错误: 文件不存在 {abs_path}"
        try:
            with open(abs_path, "r", encoding="utf-8") as f:
                json_data = json.load(f)
            return json.dumps(json_data, ensure_ascii=False, indent=indent)
        except json.JSONDecodeError as e:
            return f"错误: JSON 格式无效 - {e}"

    # ── validate: 验证格式 ──
    if operation == "validate":
        if not os.path.isfile(abs_path):
            return f"错误: 文件不存在 {abs_path}"
        try:
            with open(abs_path, "r", encoding="utf-8") as f:
                json.load(f)
            return f"JSON 格式有效: {abs_path}"
        except json.JSONDecodeError as e:
            return f"JSON 格式无效: {e}"

    # ── 以下操作需要文件存在 ──
    if not os.path.isfile(abs_path):
        return f"错误: 文件不存在 {abs_path}"

    try:
        with open(abs_path, "r", encoding="utf-8") as f:
            json_data = json.load(f)
    except json.JSONDecodeError as e:
        return f"错误: JSON 格式无效 - {e}"

    # ── write: 写入/更新键值 ──
    if operation == "write":
        if not key:
            return "错误: 请提供 key 参数"
        try:
            parsed_value = json.loads(value)
        except (json.JSONDecodeError, TypeError):
            parsed_value = value
        _set_nested(json_data, key, parsed_value)
        with open(abs_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=indent)
        return f"已设置 {key} = {json.dumps(parsed_value, ensure_ascii=False)}"

    # ── delete: 删除键 ──
    if operation == "delete":
        if not key:
            return "错误: 请提供 key 参数"
        if _delete_nested(json_data, key):
            with open(abs_path, "w", encoding="utf-8") as f:
                json.dump(json_data, f, ensure_ascii=False, indent=indent)
            return f"已删除键 {key}"
        return f"错误: 键 {key} 不存在"

    # ── merge: 合并 JSON ──
    if operation == "merge":
        if not merge_data:
            return "错误: 请提供 merge_data 参数"
        try:
            merge_obj = json.loads(merge_data)
        except json.JSONDecodeError as e:
            return f"错误: merge_data JSON 格式无效 - {e}"
        if isinstance(json_data, dict) and isinstance(merge_obj, dict):
            json_data.update(merge_obj)
        else:
            return "错误: 合并操作仅支持 JSON 对象"
        with open(abs_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=indent)
        return f"已合并 JSON 数据"

    # ── array: 向数组追加元素 ──
    if operation == "array":
        if not key:
            return "错误: 请提供 key 参数指定数组路径"
        arr = _get_nested(json_data, key)
        if arr is None:
            arr = []
            _set_nested(json_data, key, arr)
        if not isinstance(arr, list):
            return f"错误: {key} 不是数组"
        try:
            parsed_value = json.loads(data)
        except (json.JSONDecodeError, TypeError):
            parsed_value = data
        arr.append(parsed_value)
        with open(abs_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=indent)
        return f"已向 {key} 追加元素（当前长度: {len(arr)}）"

    return f"错误: 未知操作 '{operation}'，支持: create, read, write, delete, merge, array, validate"


# ══════════════════════════════════════════════════════════════════
# 己方干员查询工具
# ══════════════════════════════════════════════════════════════════

@mcp.tool()
def query_operator(name: str, max_results: int = 5) -> str:
    """
    查询明日方舟己方干员的详细数据，包括属性、天赋、技能。

    参数:
        name:         干员名称或 ID（支持模糊搜索，如 "凯尔希"、"amiya"、"Mon3tr"）
        max_results:  最大返回条数（默认 5）
    """
    _LOCAL_BASE = Path(__file__).parent / "ArknightsGameData"
    _REMOTE_BASE = "https://raw.githubusercontent.com/Kengxxiao/ArknightsGameData/master/zh_CN"

    def _data_path(rel: str) -> str:
        local = _LOCAL_BASE / rel
        if local.is_file():
            return str(local)
        return f"{_REMOTE_BASE}/{rel}"

    try:
        _chars = _load_json(_data_path("gamedata/excel/character_table.json"))
        _skills = _load_json(_data_path("gamedata/excel/skill_table.json"))
    except Exception as e:
        return f"错误: 无法加载游戏数据 {e}"

    _PROF = {"PIONEER": "先锋", "WARRIOR": "近卫", "TANK": "重装", "SNIPER": "狙击",
             "CASTER": "术师", "MEDIC": "医疗", "SUPPORTER": "辅助", "SPECIALIST": "特种"}
    _POS = {"MELEE": "近战", "RANGED": "远程"}
    _RARITY = {"TIER_1": "★", "TIER_2": "★★", "TIER_3": "★★★",
               "TIER_4": "★★★★", "TIER_5": "★★★★★", "TIER_6": "★★★★★★"}

    def _clean(s):
        return re.sub(r"<[^>]+>", "", str(s)) if s else ""

    query = name.lower().strip()
    matches = [
        (cid, d) for cid, d in _chars.items()
        if query in d.get("name", "").lower()
        or query in cid.lower()
        or query in d.get("appellation", "").lower()
    ]
    if not matches:
        return f"未找到干员 \"{name}\""

    results = []
    for cid, d in matches[:max_results]:
        lines = [f"[{cid}]"]
        lines.append(f"  名称: {d.get('name')} ({d.get('appellation', '')})")
        lines.append(f"  稀有度: {_RARITY.get(d.get('rarity'), d.get('rarity'))}")
        lines.append(f"  职业: {_PROF.get(d.get('profession'), d.get('profession'))}")
        lines.append(f"  分支: {d.get('subProfessionId')}")
        lines.append(f"  位置: {_POS.get(d.get('position'), d.get('position'))}")
        desc = _clean(d.get("description"))
        if desc:
            lines.append(f"  描述: {desc[:120]}")
        nation = d.get("nationId", "")
        if nation:
            lines.append(f"  势力: {nation}")

        # 满级属性
        phases = d.get("phases", [])
        if phases:
            attrs_list = phases[-1].get("attributesKeyFrames", [])
            if attrs_list:
                a = attrs_list[-1].get("data", {})
                lines.append("  满级属性:")
                for k, label in [("maxHp", "生命"), ("atk", "攻击"), ("def", "防御"),
                                 ("magicResistance", "法抗"), ("cost", "费用"),
                                 ("blockCnt", "阻挡"), ("baseAttackTime", "攻击间隔")]:
                    v = a.get(k)
                    if v is not None:
                        lines.append(f"    {label}: {v}")

        # 天赋
        talents = d.get("talents", [])
        if talents:
            lines.append("  天赋:")
            for t in talents:
                cands = t.get("candidates", [])
                if cands:
                    last_t = cands[-1]
                    t_name = last_t.get("name", "?")
                    t_desc = _clean(last_t.get("description", ""))[:100]
                    lines.append(f"    - {t_name}: {t_desc}")

        # 技能
        char_skills = d.get("skills", [])
        if char_skills:
            lines.append("  技能:")
            for sk in char_skills:
                sk_id = sk.get("skillId", "")
                sk_data = _skills.get(sk_id, {})
                levels = sk_data.get("levels", [])
                sk_name = levels[-1].get("name", sk_id) if levels else sk_id
                sk_desc = _clean(levels[-1].get("description", ""))[:100] if levels else ""
                lines.append(f"    - {sk_name}: {sk_desc}")

        results.append("\n".join(lines))
    return "\n\n".join(results)


# ══════════════════════════════════════════════════════════════════
# 剧情查询工具
# ══════════════════════════════════════════════════════════════════

def _parse_story_script(text: str, show_stage_direction: bool = False) -> str:
    """解析剧情脚本，提取对话和旁白"""
    lines = text.split("\n")
    result = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 对话: [name="角色"] 内容
        m = re.match(r'\[name="([^"]+)"\]\s*(.*)', line)
        if m:
            char, dialogue = m.group(1), m.group(2).strip()
            if dialogue:
                result.append(f"{char}: {dialogue}")
            continue
        # 旁白/描述行（不以 [ 开头的纯文本）
        if not line.startswith("["):
            result.append(line)
            continue
        # 可选：显示舞台指令
        if show_stage_direction:
            tag = re.match(r"\[(\w+)", line)
            if tag:
                result.append(f"[{tag.group(1)}]")
    return "\n".join(result)


@mcp.tool()
def query_story(
    query: str = "",
    story_id: str = "",
    operation: str = "search",
    max_results: int = 20,
    show_direction: bool = False,
) -> str:
    """
    查询和获取明日方舟剧情文本。

    参数:
        query:           搜索关键词（章节名、活动名、关卡代号，如 "第七章"、"骑兵与猎人"、"GT-1"）
        story_id:        剧情 ID（read 操作用，如 "obt/main/level_main_01-07_beg"）
        operation:       操作类型：
                         - search: 搜索剧情（默认）
                         - read:   读取指定剧情文本
                         - list:   列出某章节/活动的所有剧情
        max_results:     search 最大返回条数（默认 20）
        show_direction:  是否显示舞台指令（默认 false，仅显示对话）
    """
    _LOCAL_STORY = Path(__file__).parent / "ArknightsGameData"
    _REMOTE_BASE = "https://raw.githubusercontent.com/Kengxxiao/ArknightsGameData/master/zh_CN"

    def _data_path(rel: str) -> str:
        local = _LOCAL_STORY / rel
        if local.is_file():
            return str(local)
        return f"{_REMOTE_BASE}/{rel}"

    # ── 加载剧情索引 ──
    try:
        review = _load_json(_data_path("gamedata/excel/story_review_table.json"))
        story_table = _load_json(_data_path("gamedata/excel/story_table.json"))
    except Exception as e:
        return f"错误: 无法加载剧情索引 {e}"

    # story_review_table 的 storyId 格式: main_15_level_main_15-15_beg
    # story_table 的 key 格式:        obt/main/level_main_15-15_beg
    # 实际文件路径:                    gamedata/story/obt/main/level_main_15-15_beg.txt
    def _resolve_story_id(sid: str) -> str:
        """将 review_table 的 storyId 转为实际文件路径"""
        # 先直接查 story_table
        if sid in story_table:
            return sid
        # 尝试转换: main_15_level_main_15-15_beg -> obt/main/level_main_15-15_beg
        # 匹配模式: main_{chapter}_level_{rest}
        m = re.match(r"main_(\d+)_level_(.+)", sid)
        if m:
            candidate = f"obt/main/level_{m.group(2)}"
            if candidate in story_table:
                return candidate
        # 活动剧情: act3d0_level_act3d0_01_beg -> obt/activity/level_act3d0_01_beg
        m = re.match(r"act\w+_level_(.+)", sid)
        if m:
            candidate = f"obt/activity/level_{m.group(1)}"
            if candidate in story_table:
                return candidate
        return sid

    # ── search: 搜索剧情 ──
    if operation == "search":
        if not query:
            return "错误: 请提供 query 参数"
        q = query.lower().strip()
        matches = []
        for key, entry in review.items():
            name = entry.get("name", "")
            entry_type = entry.get("entryType", "")
            # 匹配活动名
            if q in name.lower() or q in key.lower():
                for iu in entry.get("infoUnlockDatas", []):
                    matches.append({
                        "activity": name,
                        "storyId": iu.get("storyId", ""),
                        "storyCode": iu.get("storyCode", ""),
                        "storyName": iu.get("storyName", ""),
                        "avgTag": iu.get("avgTag", ""),
                        "entryType": entry_type,
                    })
            else:
                # 匹配关卡代号或剧情名
                for iu in entry.get("infoUnlockDatas", []):
                    code = iu.get("storyCode") or ""
                    sname = iu.get("storyName") or ""
                    sid = iu.get("storyId") or ""
                    if q in code.lower() or q in sname.lower() or q in sid.lower():
                        matches.append({
                            "activity": name,
                            "storyId": sid,
                            "storyCode": code,
                            "storyName": sname,
                            "avgTag": iu.get("avgTag", ""),
                            "entryType": entry_type,
                        })
        if not matches:
            return f"未找到匹配 \"{query}\" 的剧情"
        lines = [f"找到 {len(matches)} 条匹配（显示前 {min(max_results, len(matches))} 条）:", ""]
        for m in matches[:max_results]:
            tag = f"[{m['avgTag']}]" if m["avgTag"] else ""
            code = m["storyCode"] or "—"
            resolved = _resolve_story_id(m["storyId"])
            lines.append(f"  {m['activity']} / {code} {m['storyName']} {tag}")
            lines.append(f"    storyId: {m['storyId']}")
            lines.append("")
        if len(matches) > max_results:
            lines.append(f"... 共 {len(matches)} 条，仅显示前 {max_results} 条")
        return "\n".join(lines)

    # ── list: 列出章节/活动剧情 ──
    if operation == "list":
        if not query:
            return "错误: 请提供 query 参数（章节名或活动名）"
        q = query.lower().strip()
        found = None
        for key, entry in review.items():
            name = entry.get("name", "")
            if q in name.lower() or q in key.lower():
                found = (key, entry)
                break
        if not found:
            return f"未找到匹配 \"{query}\" 的章节/活动"
        key, entry = found
        lines = [
            f"活动: {entry.get('name', key)}",
            f"类型: {entry.get('entryType', '?')} / {entry.get('actType', '?')}",
            f"剧情数: {len(entry.get('infoUnlockDatas', []))}",
            "",
        ]
        for iu in entry.get("infoUnlockDatas", []):
            tag = f"[{iu.get('avgTag', '')}]" if iu.get("avgTag") else ""
            code = iu.get("storyCode", "—")
            lines.append(f"  {code} {iu.get('storyName', '?')} {tag}")
            lines.append(f"    storyId: {iu.get('storyId', '')}")
        return "\n".join(lines)

    # ── read: 读取剧情文本 ──
    if operation == "read":
        if not story_id:
            return "错误: 请提供 story_id 参数"
        resolved = _resolve_story_id(story_id)
        url = _data_path(f"gamedata/story/{resolved}.txt")
        try:
            if url.startswith(("http://", "https://")):
                r = requests.get(url, timeout=15)
                if r.status_code != 200:
                    return f"错误: 无法获取剧情文件 (HTTP {r.status_code})"
                text = r.text
            else:
                with open(url, "r", encoding="utf-8") as f:
                    text = f.read()
        except Exception as e:
            return f"错误: 无法读取剧情 {e}"
        parsed = _parse_story_script(text, show_direction)
        if not parsed:
            return "该剧情文件无可提取的对话内容"
        return f"[{story_id}]\n\n{parsed}"

    return f"错误: 未知操作 '{operation}'，支持: search, read, list"


# ══════════════════════════════════════════════════════════════════
# MAA 控制工具
# ══════════════════════════════════════════════════════════════════

import ctypes
import ctypes.util
import platform
import threading
import time

_maa_lib = None
_maa_instance = None
_maa_connected = False
_maa_log_buffer: list[str] = []
_maa_log_lock = threading.Lock()

# MAA 回调
_CallBackType = ctypes.CFUNCTYPE(None, ctypes.c_int, ctypes.c_char_p, ctypes.c_void_p)


def _maa_callback(msg: int, detail: bytes, arg):
    """MAA 回调函数，收集日志"""
    try:
        detail_str = detail.decode("utf-8") if detail else ""
        with _maa_log_lock:
            _maa_log_buffer.append(f"[{msg}] {detail_str}")
            if len(_maa_log_buffer) > 200:
                _maa_log_buffer.pop(0)
    except Exception:
        pass


_callback_ref = _CallBackType(_maa_callback)


def _load_maa(maa_path: str) -> bool:
    """加载 MaaCore.dll 并初始化资源"""
    global _maa_lib
    p = Path(maa_path)
    dll_path = p / "MaaCore.dll"
    if not dll_path.exists():
        return False

    # 把 MAA 目录加入 PATH
    env_path = os.environ.get("PATH", "")
    if str(p) not in env_path:
        os.environ["PATH"] = str(p) + os.pathsep + env_path

    _maa_lib = ctypes.WinDLL(str(dll_path))
    _set_lib_properties()

    # 加载资源
    resource_path = p / "resource"
    if resource_path.exists():
        _maa_lib.AsstLoadResource(str(p).encode("utf-8"))
    return True


def _set_lib_properties():
    """设置 ctypes 接口"""
    _maa_lib.AsstSetUserDir.restype = ctypes.c_bool
    _maa_lib.AsstSetUserDir.argtypes = (ctypes.c_char_p,)
    _maa_lib.AsstLoadResource.restype = ctypes.c_bool
    _maa_lib.AsstLoadResource.argtypes = (ctypes.c_char_p,)
    _maa_lib.AsstCreate.restype = ctypes.c_void_p
    _maa_lib.AsstCreate.argtypes = ()
    _maa_lib.AsstCreateEx.restype = ctypes.c_void_p
    _maa_lib.AsstCreateEx.argtypes = (ctypes.c_void_p, ctypes.c_void_p)
    _maa_lib.AsstDestroy.argtypes = (ctypes.c_void_p,)
    _maa_lib.AsstConnect.restype = ctypes.c_bool
    _maa_lib.AsstConnect.argtypes = (ctypes.c_void_p, ctypes.c_char_p, ctypes.c_char_p, ctypes.c_char_p)
    _maa_lib.AsstAppendTask.restype = ctypes.c_int
    _maa_lib.AsstAppendTask.argtypes = (ctypes.c_void_p, ctypes.c_char_p, ctypes.c_char_p)
    _maa_lib.AsstSetTaskParams.restype = ctypes.c_bool
    _maa_lib.AsstSetTaskParams.argtypes = (ctypes.c_void_p, ctypes.c_int, ctypes.c_char_p)
    _maa_lib.AsstStart.restype = ctypes.c_bool
    _maa_lib.AsstStart.argtypes = (ctypes.c_void_p,)
    _maa_lib.AsstStop.restype = ctypes.c_bool
    _maa_lib.AsstStop.argtypes = (ctypes.c_void_p,)
    _maa_lib.AsstRunning.restype = ctypes.c_bool
    _maa_lib.AsstRunning.argtypes = (ctypes.c_void_p,)
    _maa_lib.AsstGetVersion.restype = ctypes.c_char_p


@mcp.tool()
def maa_connect(
    maa_path: str,
    adb_path: str = "",
    address: str = "127.0.0.1:5555",
    config: str = "General",
) -> str:
    """
    连接 MAA 到模拟器/设备。首次调用时加载 MaaCore.dll。

    参数:
        maa_path:  MAA 安装目录（含 MaaCore.dll 和 resource/），如 "C:/MAA"
        adb_path:  adb 路径，留空则使用 MAA 自带的 adb
        address:   设备地址，如 "127.0.0.1:5555"（MuMu）、"127.0.0.1:7555"（雷电）
        config:    连接配置，默认 "General"
    """
    global _maa_lib, _maa_instance, _maa_connected

    try:
        p = Path(maa_path)
        if not _maa_lib:
            if not _load_maa(maa_path):
                return f"错误: 在 {maa_path} 未找到 MaaCore.dll"

        # 创建实例
        if _maa_instance:
            _maa_lib.AsstStop(_maa_instance)
            _maa_lib.AsstDestroy(_maa_instance)
        _maa_instance = _maa_lib.AsstCreateEx(_callback_ref, None)

        # 确定 adb 路径
        if not adb_path:
            for candidate in ["platform-tools/adb.exe", "adb.exe"]:
                if (p / candidate).exists():
                    adb_path = str(p / candidate)
                    break
            if not adb_path:
                adb_path = "adb"

        # 连接
        ok = _maa_lib.AsstConnect(
            _maa_instance,
            adb_path.encode("utf-8"),
            address.encode("utf-8"),
            config.encode("utf-8"),
        )
        _maa_connected = bool(ok)

        if _maa_connected:
            ver = _maa_lib.AsstGetVersion().decode("utf-8")
            return f"连接成功！MAA 版本: {ver}，设备: {address}"
        else:
            return f"连接失败。请检查: 1) adb路径是否正确 2) 模拟器是否运行 3) 地址是否正确"
    except Exception as e:
        return f"错误: {e}"


@mcp.tool()
def maa_start_task(
    tasks: str,
) -> str:
    """
    添加并启动 MAA 任务。可同时添加多个任务（按顺序执行）。

    参数:
        tasks:  任务配置，JSON 数组格式。每个元素: {"type": "任务类型", "params": {参数}}
                支持的任务类型:
                - StartUp: 开始唤醒，params: {"client_type": "Official"/"Bilibili", "start_game_enabled": true}
                - Fight: 刷理智，params: {"stage": "1-7"/"CE-6"/"AP-5"/""(当前关), "times": 99, "medicine": 0}
                - Recruit: 自动公招，params: {"refresh": true, "select": [4,5], "confirm": [3,4,5]}
                - Infrast: 基建换班，params: {"facility": ["Mfg","Trade","Power","Control","Reception","Office","Dorm"]}
                - Mall: 信用购物，params: {"shopping": true}
                - Award: 领取奖励
                - Roguelike: 肉鸽，params: {"theme": "Phantom"/"Mizuki"/"Sami"/"Sarkaz"}
                - Copilot: 自动作业，params: {"filename": "作业JSON路径"}
                - CloseDown: 关闭游戏
    示例: [{"type":"StartUp","params":{"client_type":"Official"}},{"type":"Fight","params":{"stage":"1-7","times":5}}]
    """
    global _maa_instance, _maa_connected

    if not _maa_instance or not _maa_connected:
        return "错误: MAA 未连接，请先调用 maa_connect"

    try:
        # 先停止正在运行的任务
        if _maa_lib.AsstRunning(_maa_instance):
            _maa_lib.AsstStop(_maa_instance)
            time.sleep(0.5)

        task_list = json.loads(tasks)
        if not isinstance(task_list, list):
            task_list = [task_list]

        added = []
        for t in task_list:
            task_type = t.get("type", "")
            params = t.get("params", {})
            task_id = _maa_lib.AsstAppendTask(
                _maa_instance,
                task_type.encode("utf-8"),
                json.dumps(params, ensure_ascii=False).encode("utf-8"),
            )
            added.append(f"{task_type}(id={task_id})")

        ok = _maa_lib.AsstStart(_maa_instance)
        if ok:
            return f"已启动 {len(added)} 个任务: {', '.join(added)}"
        else:
            return "任务启动失败"
    except json.JSONDecodeError as e:
        return f"JSON 解析错误: {e}"
    except Exception as e:
        return f"错误: {e}"


@mcp.tool()
def maa_stop() -> str:
    """停止当前所有 MAA 任务。"""
    global _maa_instance
    if not _maa_instance:
        return "MAA 未初始化"
    _maa_lib.AsstStop(_maa_instance)
    return "已停止所有任务"


@mcp.tool()
def maa_status() -> str:
    """查询 MAA 当前状态和最近日志。"""
    global _maa_instance, _maa_connected
    if not _maa_instance:
        return "MAA 未初始化，请先调用 maa_connect"

    running = bool(_maa_lib.AsstRunning(_maa_instance))
    status = "运行中" if running else "空闲"

    with _maa_log_lock:
        recent_logs = _maa_log_buffer[-20:]

    # 解析日志提取关键信息
    info_lines = []
    for log in recent_logs:
        try:
            # 日志格式: [msg_id] json_detail
            detail = log.split("] ", 1)[1] if "] " in log else log
            d = json.loads(detail)
            if "what" in d:
                info_lines.append(d.get("what", "") + ": " + d.get("details", {}).get("task", d.get("why", "")))
            elif "details" in d and isinstance(d["details"], dict):
                sub = d["details"]
                if "task" in sub:
                    info_lines.append(sub["task"])
        except (json.JSONDecodeError, KeyError, TypeError):
            if len(log) < 200:
                info_lines.append(log)

    result = f"状态: {status}\n连接: {'已连接' if _maa_connected else '未连接'}\n"
    if info_lines:
        result += "\n最近日志:\n" + "\n".join(info_lines[-10:])
    return result


# ══════════════════════════════════════════════════════════════════
# CSGO/CS2 赛事数据工具（HLTV + Liquipedia + 5EPlay）
# ══════════════════════════════════════════════════════════════════

import time as _time

# ── cloudscraper 可选导入（惰性初始化） ──
try:
    import cloudscraper as _cloudscraper
    _HAS_CLOUDSCRAPER = True
except ImportError:
    _HAS_CLOUDSCRAPER = False

_hltv_session = None
_hltv_warmed: bool = False


def _hltv_get_session():
    """惰性获取/创建 HLTV session（延迟到首次使用时以避开 Cloudflare 时效问题）。"""
    global _hltv_session
    if _hltv_session is not None:
        return _hltv_session

    if _HAS_CLOUDSCRAPER:
        _hltv_session = _cloudscraper.create_scraper(
            browser={
                "browser": "firefox",
                "platform": "windows",
                "mobile": False,
            }
        )
        # 注意: cloudscraper 自行管理 UA/headers，覆盖会破坏 CF 绕过能力
    else:
        _hltv_session = requests.Session()
        _hltv_session.headers.update({
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/124.0.0.0 Safari/537.36"
            ),
            "Accept-Language": "en-US,en;q=0.9,zh-CN;q=0.8",
            "Referer": "https://www.hltv.org/",
        })
    return _hltv_session


def _hltv_ensure_warmup() -> None:
    """惰性预热 — 首次调用时访问 HLTV 首页以获取 cf_clearance cookie。
    若遇到 403 则尝试重建 session 后重试一次。
    """
    global _hltv_warmed, _hltv_session
    if _hltv_warmed:
        return

    sess = _hltv_get_session()

    def _try_warmup() -> bool:
        try:
            resp = sess.get("https://www.hltv.org", timeout=20)
            if resp.status_code == 200:
                _time.sleep(1.0)
                return True
        except Exception:
            pass
        return False

    if _try_warmup():
        _hltv_warmed = True
        return

    # 预热 403 — 重建 session 后重试
    _hltv_session = None
    sess = _hltv_get_session()
    if _try_warmup():
        _hltv_warmed = True
        return

    # 仍失败 — 标记已尝试，后续请求自行处理
    _hltv_warmed = True


# ── 缓存 ──
_csgo_cache: dict[str, tuple[float, str]] = {}

_CSGO_TTL = {
    "matches":     300,    # 5 分钟
    "results":     3600,   # 1 小时
    "rankings":    21600,  # 6 小时
    "tournaments": 3600,   # 1 小时
    "player":      7200,   # 2 小时
    "match":       1800,   # 30 分钟（比赛结束后数据不可变）
    "5eplay_api":  120,    # 2 分钟（5EPlay API 进行中比赛）
}


def _csgo_cache_get(key: str, ttl: int) -> Optional[str]:
    entry = _csgo_cache.get(key)
    if entry and (_time.time() - entry[0]) < ttl:
        return entry[1]
    return None


def _csgo_cache_set(key: str, data: str) -> None:
    _csgo_cache[key] = (_time.time(), data)


def _hltv_get(url: str, ttl_key: str = "matches") -> str:
    """发起 HLTV 请求，返回 HTML 文本。带缓存。"""
    global _hltv_session

    _hltv_ensure_warmup()
    cache_key = f"hltv:{url}"
    ttl = _CSGO_TTL.get(ttl_key, 300)
    cached = _csgo_cache_get(cache_key, ttl)
    if cached:
        return cached

    def _do_request():
        _time.sleep(0.5)
        return _hltv_get_session().get(url, timeout=15)

    try:
        resp = _do_request()
        # 403 — 尝试重建 session 后重试一次
        if resp.status_code == 403:
            _hltv_session = None
            _time.sleep(1.0)
            _hltv_get_session().get("https://www.hltv.org", timeout=20)
            _time.sleep(1.0)
            resp = _do_request()

        if resp.status_code == 403:
            return "ERROR:403"
        resp.raise_for_status()
        _csgo_cache_set(cache_key, resp.text)
        return resp.text
    except requests.exceptions.Timeout:
        return "ERROR:TIMEOUT"
    except Exception as e:
        return f"ERROR:{e}"


# ── 5EPlay API 封装 ──
_5EPLAY_LIST_URL = "https://app.5eplay.com/api/tournament/session_list"
_5EPLAY_MATCH_URL = "https://esports-data.5eplaycdn.com/v1/api/csgo/matches"
_5EPLAY_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/149.0.0.0 Safari/537.36"
    ),
    "Accept": "application/json",
    "Referer": "https://event.5eplay.com/",
    "Origin": "https://event.5eplay.com",
}

_5EPLAY_STATUS_MAP = {
    "upcoming": 0,
    "ongoing": 1,
    "live": 1,
    "results": 2,
}


def _5eplay_api_get(url: str, params: dict = None, ttl: int = 120) -> Optional[dict]:
    """调用 5EPlay API，返回解析后的 JSON dict。带缓存。"""
    import time as _time
    cache_key = f"5eapi:{url}:{json.dumps(params or {}, sort_keys=True)}"
    cached = _csgo_cache_get(cache_key, ttl)
    if cached:
        try:
            return json.loads(cached)
        except (json.JSONDecodeError, TypeError):
            pass

    try:
        resp = requests.get(url, headers=_5EPLAY_HEADERS, params=params, timeout=15)
        if resp.status_code != 200:
            return None
        data = resp.json()
        _csgo_cache_set(cache_key, json.dumps(data, ensure_ascii=False))
        return data
    except Exception:
        return None


def _5eplay_match_detail(match_id: str) -> Optional[dict]:
    """获取 5EPlay 比赛详情（含选手统计、veto、半场比分等）。
    返回 data.match 对象，包含 mc_info, tt_info, global_state, bouts_state 等。
    """
    url = f"{_5EPLAY_MATCH_URL}/{match_id}/data"
    data = _5eplay_api_get(url, ttl=120)
    if data and data.get("success"):
        return data.get("data", {}).get("match", {})
    return None


# ── Liquipedia API 封装 ──
_last_liquipedia_req: float = 0


def _liquipedia_request(endpoint: str, api_key: str, params: dict) -> Optional[dict]:
    """调用 Liquipedia API，自动限速 1 req/s。"""
    global _last_liquipedia_req
    elapsed = _time.time() - _last_liquipedia_req
    if elapsed < 1.0:
        _time.sleep(1.0 - elapsed)

    headers = {
        "apikey": api_key,
        "Accept": "application/json",
        "User-Agent": "Mon3tr-MCP/1.0 (https://github.com/Mon3tr-MCP)",
    }
    try:
        resp = requests.get(
            f"https://api.liquipedia.net/api/v1{endpoint}",
            headers=headers,
            params=params,
            timeout=15,
        )
        _last_liquipedia_req = _time.time()
        if resp.status_code == 200:
            return resp.json()
        return None
    except Exception:
        return None


# ── 比赛详情解析（match detail scraping） ──


def _parse_stats_table(table) -> list[dict]:
    """解析单张 HLTV totalstats 表格，返回选手数据列表。

    预期列：PlayerName | K-D | eK-eD | Swing | ADR | eADR | KAST | eKAST | Rating 3.0
    """
    players = []
    rows = table.select("tbody tr, tr")
    for row in rows:
        cells = row.select("td, th")
        # 至少需要 选手名 + K-D + ADR + Rating 4 列有意义数据
        if len(cells) < 5:
            continue
        name = cells[0].get_text(" ", strip=True)
        if not name or len(name) < 2:
            continue
        # 跳过标题行
        name_lower = name.lower()
        if name_lower in ("player", "name", "team", "legacy", "tyloo"):
            continue

        player = {
            "name": name,
            "kd": cells[1].get_text(strip=True) if len(cells) > 1 else "",
            "adr": cells[4].get_text(strip=True) if len(cells) > 4 else "",
            "kast": cells[6].get_text(strip=True) if len(cells) > 6 else "",
            "rating": cells[8].get_text(strip=True) if len(cells) > 8 else "",
        }
        players.append(player)
    return players


def _scrape_map_scores(soup) -> list[dict]:
    """从 HLTV 比赛详情页提取每张地图的比分和 veto 信息。

    解析 .mapholder 元素，提取地图名、双方比分、半场比分、胜负标记。

    返回结构：
        [{"map": "Mirage", "status": "played",
          "team1": {"name": "Legacy", "score": "13", "result": "won"},
          "team2": {"name": "TYLOO", "score": "7", "result": "lost"},
          "half_scores": "(8:4;5:3)", "mapstats_id": "231176"}, ...]
    """
    map_scores: list[dict] = []
    for mh in soup.select(".mapholder"):
        map_name_el = mh.select_one(".mapname")
        map_name = map_name_el.get_text(strip=True) if map_name_el else "Unknown"

        # 判断是否已打 (played vs optional)
        status_el = mh.select_one(".played, .optional")
        status_classes = status_el.get("class", []) if status_el else []
        status = "played" if "played" in status_classes else "optional"

        team1: dict = {"name": "", "score": "", "result": ""}
        team2: dict = {"name": "", "score": "", "result": ""}
        half_scores = ""
        mapstats_id = ""

        results = mh.select_one(".results")
        if results:
            # 提取 mapstats_id
            stats_link = results.select_one("a.results-stats")
            if stats_link:
                href = stats_link.get("href", "")
                m = re.search(r"/mapstatsid/(\d+)/", href)
                if m:
                    mapstats_id = m.group(1)

            # 提取半场比分
            half_el = results.select_one(".results-center-half-score")
            if half_el:
                half_scores = half_el.get_text(strip=True)

            # 提取队伍1 (左侧)
            left = results.select_one(".results-left")
            if left:
                left_classes = left.get("class", [])
                if "won" in left_classes:
                    team1["result"] = "won"
                elif "lost" in left_classes:
                    team1["result"] = "lost"
                else:
                    team1["result"] = "tie"
                name_el = left.select_one(".results-teamname")
                score_el = left.select_one(".results-team-score")
                team1["name"] = name_el.get_text(strip=True) if name_el else ""
                team1["score"] = score_el.get_text(strip=True) if score_el else ""

            # 提取队伍2 (右侧)
            right = results.select_one(".results-right")
            if right:
                right_classes = right.get("class", [])
                if "won" in right_classes:
                    team2["result"] = "won"
                elif "lost" in right_classes:
                    team2["result"] = "lost"
                else:
                    team2["result"] = "tie"
                name_el = right.select_one(".results-teamname")
                score_el = right.select_one(".results-team-score")
                team2["name"] = name_el.get_text(strip=True) if name_el else ""
                team2["score"] = score_el.get_text(strip=True) if score_el else ""

        map_scores.append({
            "map": map_name,
            "status": status,
            "team1": team1,
            "team2": team2,
            "half_scores": half_scores,
            "mapstats_id": mapstats_id,
        })

    return map_scores


def _scrape_veto_info(soup) -> str:
    """提取比赛 veto/选图信息。"""
    veto_boxes = soup.select(".veto-box")
    if not veto_boxes:
        return ""

    parts: list[str] = []
    for box in veto_boxes:
        text = box.get_text(" ", strip=True)
        text = re.sub(r"\s{2,}", " ", text)
        if text and len(text) > 2:
            parts.append(text)

    # 按序号分行：查找编号列表和普通文本
    result = " ".join(parts)
    # 在序号前插入换行: "2. TYLOO" -> "\n2. TYLOO"
    result = re.sub(r"(?<=\S)\s*(\d+\.\s)", r"\n\1", result)
    return result


def _scrape_match_stats(soup, map_scores: Optional[list] = None) -> list[dict]:
    """从 HLTV 比赛详情页 soup 提取每张地图的选手统计。

    参数:
        soup:       比赛详情页的 BeautifulSoup 对象
        map_scores: 可选，_scrape_map_scores() 的返回值，用于将 stats-content
                    div id 映射到地图名

    返回结构：
        [{"map": "Overall"|"Ancient"|..., "team": "Legacy"|"TYLOO", "players": [...]}, ...]
    """
    # 构建 mapstats_id → 地图名 的映射
    id_to_map: dict = {}
    if map_scores:
        for ms in map_scores:
            mid = ms.get("mapstats_id", "")
            if mid:
                id_to_map[mid] = ms["map"]

    all_stats: list[dict] = []

    # 遍历所有 .stats-content div（#all-content 及每个地图的 #{mapstatsid}-content）
    for content_div in soup.select(".stats-content"):
        div_id = content_div.get("id", "")

        # 确定地图名
        if div_id == "all-content":
            map_name = "Overall"
        elif div_id.endswith("-content"):
            stats_id = div_id.replace("-content", "")
            map_name = id_to_map.get(stats_id, f"Map({stats_id})")
        else:
            continue

        tables = content_div.select("table.totalstats")
        for table in tables:
            team_name = ""
            header_cell = table.select_one("tr > th, tr > td")
            if header_cell:
                team_name = header_cell.get_text(strip=True)
            players = _parse_stats_table(table)
            if players:
                all_stats.append({
                    "map": map_name,
                    "team": team_name,
                    "players": players,
                })

    return all_stats


# ── 工具 1: csgo_matches ──

@mcp.tool()
def csgo_matches(
    match_type: str = "upcoming",
    source: str = "hltv",
    offset: int = 0,
    limit: int = 20,
) -> str:
    """
    获取 CSGO/CS2 比赛信息。

    参数:
        match_type: 比赛类型 — "upcoming"(即将开始), "results"(已结束), "live"(正在进行)
        source:     数据源 — "hltv"(HLTV,默认), "5eplay"(5EPlay,中国/亚洲赛事)
        offset:     分页偏移量（HLTV results 每页 100 条）
        limit:      最大返回条数（默认 20，上限 50）

    返回:
        比赛列表，包含队伍名、比分/时间、赛事名、比赛链接
    """
    limit = min(max(limit, 1), 50)

    if source == "5eplay":
        return _csgo_matches_5eplay(match_type, limit)

    # ── HLTV 数据源 ──
    if match_type == "upcoming":
        url = "https://www.hltv.org/matches"
        ttl_key = "matches"
    elif match_type == "results":
        url = f"https://www.hltv.org/results?offset={offset}"
        ttl_key = "results"
    elif match_type == "live":
        # HLTV 的 live 标记依赖 JS 动态渲染，HTML 抓取不可靠
        return (
            "⚠️ HLTV live 模式不可靠（依赖动态渲染）。\n"
            "请使用 source=\"5eplay\" 获取实时比赛:\n"
            "  csgo_matches(match_type=\"live\", source=\"5eplay\")\n\n"
            + _csgo_matches_5eplay("live", limit)
        )
    else:
        return f"错误: 无效的 match_type '{match_type}'，支持: upcoming, results, live"

    raw = _hltv_get(url, ttl_key)
    if raw.startswith("ERROR:"):
        if "403" in raw:
            return "错误: HLTV 请求被 Cloudflare 拦截，请运行 pip install cloudscraper 安装绕过库"
        return f"错误: HLTV 请求失败 {raw}"

    soup = BeautifulSoup(raw, "html.parser")
    results = []

    if match_type in ("upcoming", "live"):
        # 解析即将开始 / 正在进行的比赛
        match_containers = soup.select(".upcoming-match")
        if not match_containers:
            # 备用选择器
            match_containers = soup.select("[class*='match']")

        for mc in match_containers:
            team_els = mc.select(".team-name, .match-teamname .team")
            if len(team_els) < 2:
                continue
            team1 = team_els[0].get_text(strip=True)
            team2 = team_els[1].get_text(strip=True)

            time_el = mc.select_one(".match-time, .matchInfo .time, [class*='time']")
            match_time = time_el.get_text(strip=True) if time_el else "未知"

            event_el = mc.select_one(".match-event, .event-name, [class*='event']")
            event_name = event_el.get_text(strip=True) if event_el else ""

            format_el = mc.select_one(".match-format, [class*='format']")
            match_format = format_el.get_text(strip=True) if format_el else ""

            link_el = mc.select_one("a[href*='/matches/']")
            link = ""
            if link_el:
                href = link_el.get("href", "")
                link = f"https://www.hltv.org{href}" if href.startswith("/") else href

            line = f"{team1} vs {team2}"
            if match_format:
                line += f" ({match_format})"
            if event_name:
                line += f"\n  赛事: {event_name}"
            line += f"\n  时间: {match_time}"
            if link:
                line += f"\n  链接: {link}"
            results.append(line)

    elif match_type == "results":
        # 解析已结束的比赛结果
        result_containers = soup.select(".result-con")
        if not result_containers:
            result_containers = soup.select("[class*='result']")

        for rc in result_containers:
            team_els = rc.select(".team-name, .team")
            score_els = rc.select(".score")
            if len(team_els) < 2:
                continue

            team1 = team_els[0].get_text(strip=True)
            team2 = team_els[1].get_text(strip=True)
            score1 = score_els[0].get_text(strip=True) if len(score_els) > 0 else "?"
            score2 = score_els[1].get_text(strip=True) if len(score_els) > 1 else "?"

            event_el = rc.select_one(".event-name, [class*='event']")
            event_name = event_el.get_text(strip=True) if event_el else ""

            map_el = rc.select_one(".map-text, [class*='map']")
            map_played = map_el.get_text(strip=True) if map_el else ""

            link_el = rc.select_one("a[href*='/matches/']")
            link = ""
            if link_el:
                href = link_el.get("href", "")
                link = f"https://www.hltv.org{href}" if href.startswith("/") else href

            line = f"{team1} {score1} - {score2} {team2}"
            if event_name:
                line += f"  [{event_name}]"
            if map_played:
                line += f"  ({map_played})"
            if link:
                line += f"\n  链接: {link}"
            results.append(line)

    if not results:
        return f"未找到{match_type}比赛数据（HLTV 页面结构可能已变化）"

    header = {
        "upcoming": "HLTV 即将开始的比赛",
        "results": "HLTV 最近比赛结果",
        "live": "HLTV 正在进行的比赛",
    }[match_type]

    output = f"{header}（共 {len(results[:limit])} 场）\n" + "=" * 50 + "\n"
    output += "\n\n".join(results[:limit])
    return output


def _csgo_matches_5eplay(match_type: str, limit: int) -> str:
    """从 5EPlay API 获取比赛数据（含每图比分）。"""
    gs = _5EPLAY_STATUS_MAP.get(match_type)
    if gs is None:
        return f"错误: 无效的 match_type '{match_type}'"

    data = _5eplay_api_get(_5EPLAY_LIST_URL, {
        "game_status": gs, "game_type": 1, "grades": "", "page": 1, "limit": limit,
    }, ttl=120 if match_type == "ongoing" else 300)

    if not data or not data.get("success"):
        return "错误: 5EPlay API 请求失败"

    matches = data.get("data", {}).get("matches", [])
    if not matches:
        return f"5EPlay 未找到{match_type}比赛"

    results: list[str] = []
    type_name = {"upcoming": "即将开始", "ongoing": "进行中", "live": "LIVE", "results": "已结束"}[match_type]
    is_live = match_type == "live"

    for m in matches:
        info = m.get("mc_info", {})
        state = m.get("state", {})
        t1 = info.get("t1_info", {})
        t2 = info.get("t2_info", {})

        t1n = t1.get("disp_name", "TBD")
        t2n = t2.get("disp_name", "TBD")

        line = f"{t1n} vs {t2n}"

        # 赛事 + 轮次
        stage = info.get("tt_stage_desc", "") or info.get("round_name", "")
        if stage:
            line += f"\n  赛事: {stage}"
        line += f"  (BO{info.get('format', '?')})"

        if match_type != "upcoming":
            bouts = state.get("bout_states", [])
            for b in bouts:
                if b.get("display") != "1":
                    continue
                mname = b.get("map_name", "?")
                s1 = b.get("t1_score", "?")
                s2 = b.get("t2_score", "?")
                status = b.get("status", "0")

                if status == "2":
                    # 已完成的地图
                    fh = f"{b.get('t1_fh_rounds','?')}:{b.get('t2_fh_rounds','?')}"
                    sh = f"{b.get('t1_sc_rounds','?')}:{b.get('t2_sc_rounds','?')}"
                    line += f"\n    {mname}: {s1}-{s2} (FH {fh}, SH {sh})"
                elif status == "1":
                    # 正在打的地图 🔴
                    rn = b.get("round_num", "?")
                    line += f"\n    🔴 {mname}: {s1}-{s2} (第{rn}回合)"

            # 总分
            qs1 = state.get("t1_quick_score", "")
            qs2 = state.get("t2_quick_score", "")
            if qs1 or qs2:
                line += f"\n  总比分: {qs1}-{qs2}"

        link = f"  详情ID: {info.get('id', '')}"
        results.append(f"{line}\n{link}")

    header = f"5EPlay {type_name}比赛（共 {len(results)} 场）"
    return f"{header}\n{'=' * 50}\n" + "\n\n".join(results)


# ── 工具 2: csgo_rankings ──

@mcp.tool()
def csgo_rankings(region: str = "world", source: str = "hltv") -> str:
    """
    获取 CSGO/CS2 战队世界排名。

    参数:
        region: 排名区域 — "world"(全球), "eu"(欧洲), "na"(北美),
                "asia"(亚洲), "sa"(南美), "oce"(大洋洲), "cis"(独联体)
        source: 数据源 — "hltv"(默认), "5eplay"(5EPlay 亚洲排名)

    返回:
        战队排名列表，包含排名、战队名、积分、阵容
    """
    if source == "5eplay":
        return _csgo_rankings_5eplay()

    # ── HLTV 排名 ──
    valid_regions = {"world", "eu", "na", "asia", "sa", "oce", "cis"}
    if region not in valid_regions:
        return f"错误: 无效的区域 '{region}'，支持: {', '.join(sorted(valid_regions))}"

    url = "https://www.hltv.org/ranking/teams/"
    if region != "world":
        url += region

    raw = _hltv_get(url, "rankings")
    if raw.startswith("ERROR:"):
        if "403" in raw:
            return "错误: HLTV 请求被 Cloudflare 拦截，请运行 pip install cloudscraper 安装绕过库"
        return f"错误: HLTV 请求失败 {raw}"

    soup = BeautifulSoup(raw, "html.parser")
    ranked_teams = soup.select(".ranked-team, [class*='ranked-team']")

    if not ranked_teams:
        return "未找到排名数据（HLTV 页面结构可能已变化）"

    region_name = {
        "world": "全球", "eu": "欧洲", "na": "北美",
        "asia": "亚洲", "sa": "南美", "oce": "大洋洲", "cis": "独联体",
    }[region]

    lines = [f"HLTV {region_name}战队排名（前 {min(30, len(ranked_teams))} 名）", "=" * 50, ""]

    for i, team in enumerate(ranked_teams[:30]):
        # 排名
        pos_el = team.select_one(".position, [class*='position']")
        rank = pos_el.get_text(strip=True) if pos_el else str(i + 1)

        # 战队名
        name_el = team.select_one(".name, .team-name, [class*='name']")
        name = name_el.get_text(strip=True) if name_el else "未知"

        # 积分
        points_el = team.select_one(".points, [class*='points']")
        points = points_el.get_text(strip=True) if points_el else ""

        # 阵容
        players = []
        for nick in team.select(".ranking-nick, [class*='nick'], [class*='player']"):
            p = nick.get_text(strip=True)
            if p:
                players.append(p)

        line = f"  {rank}. {name}"
        if points:
            line += f"  ({points}分)"
        if players:
            line += f"\n     阵容: {', '.join(players)}"
        lines.append(line)

    return "\n".join(lines)


def _csgo_rankings_5eplay() -> str:
    """从 5EPlay API 提取战队排名（从比赛数据中获取队伍排名信息）。"""
    # 尝试从正在进行的比赛中提取队伍排名
    data = _5eplay_api_get(_5EPLAY_LIST_URL, {
        "game_status": 1, "game_type": 1, "grades": "", "page": 1, "limit": 10,
    }, ttl=300)

    if not data or not data.get("success"):
        return "错误: 5EPlay API 请求失败"

    matches = data.get("data", {}).get("matches", [])
    teams: dict[str, dict] = {}
    for m in matches:
        info = m.get("mc_info", {})
        for side in ("t1_info", "t2_info"):
            t = info.get(side, {})
            tid = t.get("id", "")
            if tid and tid not in teams:
                rank = t.get("rank", "")
                if rank and str(rank).isdigit():
                    teams[tid] = {
                        "name": t.get("disp_name", "?"),
                        "rank": int(rank),
                        "logo": t.get("logo", ""),
                    }

    if not teams:
        return "5EPlay 暂未提供战队排名 API，建议使用 source=\"hltv\" 获取排名"

    ranked = sorted(teams.values(), key=lambda x: x["rank"])[:30]
    lines = [f"5EPlay 战队排名（从赛事数据提取，共 {len(ranked)} 支）", "=" * 50, ""]
    for t in ranked:
        lines.append(f"  #{t['rank']}  {t['name']}")
    return "\n".join(lines)


# ── 工具 3: csgo_tournaments ──

@mcp.tool()
def csgo_tournaments(
    status: str = "ongoing",
    source: str = "hltv",
    liquipedia_api_key: str = "",
) -> str:
    """
    获取 CSGO/CS2 赛事信息。

    参数:
        status:            赛事状态 — "ongoing"(进行中), "upcoming"(即将开始), "completed"(已结束)
        source:            数据源 — "hltv"(默认), "5eplay"(5EPlay), "liquipedia"(仅 Liquipedia)
        liquipedia_api_key: Liquipedia API Key（可选，提供后会补充赛事详情）

    返回:
        赛事列表，包含赛事名、日期、奖金池、参赛队伍数
    """
    if source == "5eplay":
        return _csgo_tournaments_5eplay(status)
    if source == "liquipedia":
        if not liquipedia_api_key:
            return "错误: 使用 Liquipedia 数据源需要提供 liquipedia_api_key 参数\n请在 https://liquipedia.net/counterstrike/Special:ApiAccount 注册获取"
        return _csgo_tournaments_liquipedia(status, liquipedia_api_key)

    # ── HLTV 赛事 ──
    if status == "ongoing":
        url = "https://www.hltv.org/events"
    elif status == "upcoming":
        url = "https://www.hltv.org/events"
    elif status == "completed":
        url = "https://www.hltv.org/events/archive"
    else:
        return f"错误: 无效的 status '{status}'，支持: ongoing, upcoming, completed"

    raw = _hltv_get(url, "tournaments")
    if raw.startswith("ERROR:"):
        if "403" in raw:
            return "错误: HLTV 请求被 Cloudflare 拦截，请运行 pip install cloudscraper 安装绕过库"
        return f"错误: HLTV 请求失败 {raw}"

    soup = BeautifulSoup(raw, "html.parser")
    events = soup.select(
        ".events-page-event, .event-item, [class*='event-card'], "
        "[class*='event-item'], .big-event, .small-event"
    )

    results = []
    for ev in events:
        name_el = ev.select_one(
            "[class*='event-name'], [class*='name'], h3, h4, a"
        )
        name = name_el.get_text(strip=True) if name_el else ""
        if not name:
            continue

        date_el = ev.select_one("[class*='date'], [class*='time'], time")
        date = date_el.get_text(strip=True) if date_el else ""

        prize_el = ev.select_one("[class*='prize'], [class*='prizepool']")
        prize = prize_el.get_text(strip=True) if prize_el else ""

        teams_el = ev.select_one("[class*='teams'], [class*='participant']")
        teams_count = teams_el.get_text(strip=True) if teams_el else ""

        location_el = ev.select_one("[class*='location'], [class*='country']")
        location = location_el.get_text(strip=True) if location_el else ""

        link_el = ev.select_one("a[href*='/event/'], a[href*='/events/']")
        link = ""
        if link_el:
            href = link_el.get("href", "")
            link = f"https://www.hltv.org{href}" if href.startswith("/") else href

        line = f"  {name}"
        if date:
            line += f"\n    日期: {date}"
        if location:
            line += f"\n    地点: {location}"
        if prize:
            line += f"\n    奖金: {prize}"
        if teams_count:
            line += f"\n    队伍: {teams_count}"
        if link:
            line += f"\n    链接: {link}"
    # ── 按状态过滤 ──
    if status == "ongoing":
        results = [
            r for r, ev in zip(results, events)
            if ev.select_one("[class*='live'], [class*='ongoing'], .event-live")
        ]
    elif status == "upcoming":
        results = [
            r for r, ev in zip(results, events)
            if not ev.select_one("[class*='live'], [class*='ongoing'], .event-live")
        ]
    # else: "completed" — 保留全部结果

    # 可选: 补充 Liquipedia 数据
    if liquipedia_api_key and results:
        lp_data = _liquipedia_request(
            "/tournaments",
            liquipedia_api_key,
            {"wiki": "counterstrike", "limit": 10, "order": "desc"},
        )
        if lp_data and "result" in lp_data:
            lp_names = {t.get("name", "").lower(): t for t in lp_data["result"]}
            for i, line in enumerate(results):
                for lp_name, lp_info in lp_names.items():
                    if lp_name and lp_name in line.lower():
                        extras = []
                        if lp_info.get("prizepool"):
                            extras.append(f"奖金(Liquipedia): {lp_info['prizepool']}")
                        if lp_info.get("format"):
                            extras.append(f"赛制: {lp_info['format']}")
                        if extras:
                            results[i] += "\n    " + "\n    ".join(extras)

    if not results:
        return f"未找到{status}赛事数据（HLTV 页面结构可能已变化）"

    status_name = {"ongoing": "进行中", "upcoming": "即将开始", "completed": "已结束"}[status]
    output = f"HLTV {status_name}赛事（共 {len(results)} 个）\n" + "=" * 50 + "\n"
    output += "\n\n".join(results[:20])
    return output


def _csgo_tournaments_5eplay(status: str) -> str:
    """从 5EPlay API 获取赛事信息。"""
    gs_map = {"ongoing": 1, "upcoming": 0, "completed": 2}
    gs = gs_map.get(status, 1)

    data = _5eplay_api_get(_5EPLAY_LIST_URL, {
        "game_status": gs, "game_type": 1, "grades": "", "page": 1, "limit": 10,
    }, ttl=300)

    if not data or not data.get("success"):
        return "错误: 5EPlay API 请求失败"

    matches = data.get("data", {}).get("matches", [])
    if not matches:
        return f"5EPlay 未找到{status}赛事"

    # 从比赛数据中聚合赛事信息
    tournaments: dict[str, dict] = {}
    for m in matches:
        info = m.get("mc_info", {})
        tt_stage = info.get("tt_stage_desc", "") or info.get("tt_stage", "")
        if tt_stage and tt_stage not in tournaments:
            t1 = info.get("t1_info", {})
            t2 = info.get("t2_info", {})
            tournaments[tt_stage] = {
                "stage": tt_stage,
                "format": f"BO{info.get('format', '?')}",
                "teams": f"{t1.get('disp_name', '?')} vs {t2.get('disp_name', '?')}",
                "round": info.get("round_name", ""),
            }

    if not tournaments:
        return f"5EPlay 未找到{status}赛事，建议使用 source=\"hltv\""

    status_name = {"ongoing": "进行中", "upcoming": "即将开始", "completed": "已结束"}[status]
    lines = [f"5EPlay {status_name}赛事（共 {len(tournaments)} 个）", "=" * 50, ""]
    for t in list(tournaments.values())[:20]:
        lines.append(f"  {t['stage']} ({t['format']})")
        lines.append(f"    对阵: {t['teams']}")
        if t["round"]:
            lines.append(f"    轮次: {t['round']}")
    return "\n".join(lines)


def _csgo_tournaments_liquipedia(status: str, api_key: str) -> str:
    """从 Liquipedia API 获取赛事信息。"""
    status_map = {
        "ongoing": "ongoing",
        "upcoming": "upcoming",
        "completed": "completed",
    }
    lp_status = status_map.get(status, "")

    params = {"wiki": "counterstrike", "limit": 20}
    if lp_status:
        params["conditions"] = f"[[status::{lp_status}]]"

    data = _liquipedia_request("/tournaments", api_key, params)
    if not data or "result" not in data:
        return "错误: Liquipedia API 请求失败，请检查 API Key 是否有效"

    tournaments = data["result"]
    if not tournaments:
        return f"未找到 Liquipedia {status} 赛事"

    status_name = {"ongoing": "进行中", "upcoming": "即将开始", "completed": "已结束"}[status]
    lines = [f"Liquipedia {status_name}赛事（共 {len(tournaments)} 个）", "=" * 50, ""]

    for t in tournaments:
        name = t.get("name", "未知")
        line = f"  {name}"
        if t.get("dates"):
            line += f"\n    日期: {t['dates']}"
        if t.get("location"):
            line += f"\n    地点: {t['location']}"
        if t.get("prizepool"):
            line += f"\n    奖金: {t['prizepool']}"
        if t.get("format"):
            line += f"\n    赛制: {t['format']}"
        if t.get("participantcount"):
            line += f"\n    参赛队伍: {t['participantcount']}"
        if t.get("pagename"):
            line += f"\n    链接: https://liquipedia.net/counterstrike/{t['pagename']}"
        lines.append(line)

    return "\n".join(lines)


# ── 工具 4: csgo_match_detail ──

@mcp.tool()
def csgo_match_detail(match_id: str, source: str = "hltv") -> str:
    """
    获取比赛详情，包含每张地图的比分、半场比分、veto 信息及选手统计数据。

    参数:
        match_id: 比赛 ID
                  HLTV 格式: "2394896/legacy-vs-tyloo-iem-cologne-major-2026-stage-2"
                  5EPlay 格式: "csgo_mc_2394896" 或纯数字 "2394896"
        source:   数据源 — "hltv"(默认), "5eplay"

    返回:
        比赛总览（总比分）、地图选择（veto）、每张地图比分（含半场）、
        每张地图双方选手统计数据（Rating、K-D、ADR、KAST、HS% 等）
    """
    if source == "5eplay":
        return _csgo_match_detail_5eplay(match_id)

    if not match_id or "/" not in match_id:
        return "错误: match_id 格式无效，例: 2394896/legacy-vs-tyloo-iem-cologne-major-2026-stage-2"

    url = f"https://www.hltv.org/matches/{match_id}"
    raw = _hltv_get(url, "match")
    if raw.startswith("ERROR:"):
        if "403" in raw:
            return "错误: HLTV 请求被 Cloudflare 拦截，请运行 pip install cloudscraper 安装绕过库"
        return f"错误: 无法获取比赛 {match_id} 的数据 ({raw})"

    soup = BeautifulSoup(raw, "html.parser")

    # 1) 提取地图比分
    map_scores = _scrape_map_scores(soup)
    if not map_scores:
        return (
            f"错误: 无法解析比赛 {match_id} 的地图数据。\n"
            "可能原因: 比赛ID不存在、HLTV 请求被拦截、或页面结构变化。"
        )

    # 2) 提取 veto 信息
    veto_info = _scrape_veto_info(soup)

    # 3) 提取选手统计
    stats = _scrape_match_stats(soup, map_scores)

    # ═══ 格式化输出 ═══

    # ── 提取比赛元信息 ──
    # 从 <title> 提取赛事名
    title_el = soup.select_one("title")
    event_name = ""
    if title_el:
        title_text = title_el.get_text(strip=True)
        # 格式: "Legacy vs. TYLOO at IEM Cologne Major 2026 Stage 2 | HLTV.org"
        m = re.search(r"\s+at\s+(.+?)\s*\|", title_text)
        if m:
            event_name = m.group(1).strip()

    # 整体比分
    score_el = soup.select_one(".score")
    overall_score = score_el.get_text(strip=True) if score_el else ""

    output = f"HLTV 比赛详情\n{'=' * 60}\n"

    # ── 比赛概要 ──
    played_maps = [ms for ms in map_scores if ms["status"] == "played"]
    team1_name = played_maps[0]["team1"]["name"] if played_maps else (map_scores[0]["team1"]["name"] if map_scores else "Team1")
    team2_name = played_maps[0]["team2"]["name"] if played_maps else (map_scores[0]["team2"]["name"] if map_scores else "Team2")

    output += f"\n⚔️  {team1_name}  vs  {team2_name}\n"
    if event_name:
        output += f"📡 赛事: {event_name}\n"
    if overall_score:
        output += f"📊 总比分: {overall_score}\n"

    if played_maps:
        team1_wins = sum(1 for ms in played_maps if ms["team1"]["result"] == "won")
        team2_wins = sum(1 for ms in played_maps if ms["team2"]["result"] == "won")
        output += f"🎯 最终: {team1_name} [{team1_wins}] — [{team2_wins}] {team2_name}"
        output += f"  (BO{len(map_scores)}, 已打 {len(played_maps)} 张)\n"

    # ── 地图选择 (Veto) ──
    if veto_info:
        output += f"\n📋 地图选择 (Veto)\n{'-' * 40}\n"
        veto_lines = veto_info.split("\n")
        for line in veto_lines:
            line = line.strip()
            if line:
                # 标注格式行
                if re.match(r"^(Best of|BO\d)", line, re.IGNORECASE):
                    output += f"  📌 {line}\n"
                elif re.match(r"^\d+\.", line):
                    output += f"     {line}\n"
                else:
                    output += f"  {line}\n"

    # ── 每张地图比分 ──
    if played_maps:
        output += f"\n🗺️  地图比分\n{'-' * 40}\n"
        for ms in map_scores:
            if ms["status"] == "played":
                t1 = ms["team1"]
                t2 = ms["team2"]
                winner_mark = "👑" if t1["result"] == "won" else ("👑" if t2["result"] == "won" else "")
                output += (
                    f"  {ms['map']:<12s} "
                    f"{t1['name']} {t1['score']} — {t2['score']} {t2['name']}"
                )
                if ms["half_scores"]:
                    output += f"    半场: {ms['half_scores']}"
                output += "\n"
            else:
                output += f"  {ms['map']:<12s} (未进行)\n"

    # ── 选手统计 ──
    if stats:
        output += f"\n📊 选手统计 (Rating 3.0)\n{'=' * 60}\n"
        for block in stats:
            map_name = block["map"]
            team = block.get("team", "")
            players = block.get("players", [])

            header = f"\n📊 {map_name}"
            if team:
                header += f" — {team}"
            output += header + f"\n{'-' * 40}\n"

            for p in players:
                output += (
                    f"  {p['name']:<28s} "
                    f"K-D: {p['kd']:<9s} "
                    f"ADR: {p['adr']:<7s} "
                    f"KAST: {p['kast']:<7s} "
                    f"Rating: {p['rating']}\n"
                )
    else:
        output += "\n⚠️ 未找到选手统计数据（页面可能未包含 stats 表格）\n"

    return output


# ── 选手搜索 / 资料页抓取 ──


def _search_hltv_player(query: str) -> list[tuple[str, str]]:
    """在 HLTV 搜索选手，返回 [(显示名, player_path), ...]。
    player_path 形如 "20702/jee"。
    """
    _hltv_ensure_warmup()
    try:
        _time.sleep(0.3)
        resp = _hltv_get_session().get(
            f"https://www.hltv.org/search?query={query}",
            timeout=10,
            headers={
                "X-Requested-With": "XMLHttpRequest",
                "Accept": "text/html, */*",
                "Referer": "https://www.hltv.org/",
            },
        )
        if resp.status_code != 200:
            return []
    except Exception:
        return []

    soup = BeautifulSoup(resp.text, "html.parser")
    results: list[tuple[str, str]] = []
    seen = set()
    for a in soup.select("a[href*='/player/']"):
        href = a.get("href", "")
        name = a.get_text(strip=True)
        # 过滤 Faceit 等无关链接
        if "#tab-faceit" in href or not name or len(name) < 2:
            continue
        # 提取纯路径 "20702/jee"
        path = href.rstrip("/").split("/player/")[-1].split("#")[0]
        if path not in seen:
            seen.add(path)
            results.append((name, path))
    return results


def _scrape_player_profile(player_path: str) -> dict:
    """抓取 HLTV 选手资料页，提取统计摘要和近期比赛。

    返回:
        {"name": ..., "team": ..., "rating": ..., "stats": {...}, "recent_matches": [...]}
    """
    url = f"https://www.hltv.org/player/{player_path}"
    raw = _hltv_get(url, "player")
    if raw.startswith("ERROR:"):
        return {}

    soup = BeautifulSoup(raw, "html.parser")

    result: dict = {"name": "", "team": "", "rating": "", "stats": {}, "recent_matches": []}

    # 选手名
    name_el = soup.select_one(".playerNickname, h1")
    if name_el:
        result["name"] = name_el.get_text(strip=True)

    # 战队
    team_el = soup.select_one(".playerTeam a, .playerTeam, [class*='team'] a")
    if team_el:
        result["team"] = team_el.get_text(strip=True)

    # Rating 3.0（近3个月）
    for stat_div in soup.select(".player-stat"):
        label_el = stat_div.select_one("b")
        value_el = stat_div.select_one(".statsVal p, .statsVal")
        if label_el and value_el:
            label = label_el.get_text(strip=True)
            value = value_el.get_text(strip=True)
            if "Rating" in label:
                result["rating"] = value
            result["stats"][label] = value

    # 近期比赛
    results_section = soup.find("h2", string=lambda t: t and "latest results" in t.lower())
    if results_section:
        table = results_section.find_next("table")
        if table:
            for row in table.select("tbody tr, tr"):
                cells = row.select("td, th")
                if len(cells) < 2:
                    continue
                texts = [c.get_text(" ", strip=True) for c in cells]
                if texts and any(t for t in texts if t):
                    result["recent_matches"].append(texts)

    return result


# ── 工具 5: csgo_player_stats ──

@mcp.tool()
def csgo_player_stats(
    query: str = "",
    stat_type: str = "rating",
    maps_filter: int = 0,
    source: str = "hltv",
    player_id: str = "",
    match_id: str = "",
) -> str:
    """
    获取 CSGO/CS2 选手统计数据。

    参数:
        query:       选手名搜索（为空则返回排行榜）
        stat_type:   统计类型 — "rating"(Rating 2.0), "kills"(击杀), "deaths"(死亡),
                     "adr"(ADR), "headshot"(爆头率), "clutch"(残局胜率)
        maps_filter: 最低场次过滤（0 = 使用 HLTV 默认值）
        source:      数据源 — "hltv"(默认), "5eplay"(5EPlay)

    返回:
        选手统计列表
    """
    if source == "5eplay":
        if match_id:
            return _csgo_player_stats_5eplay_match(match_id, stat_type)
        return _csgo_player_stats_5eplay(query, stat_type)

    valid_stats = {"rating", "kills", "deaths", "adr", "headshot", "clutch"}
    if stat_type not in valid_stats:
        return f"错误: 无效的 stat_type '{stat_type}'，支持: {', '.join(sorted(valid_stats))}"

    stat_labels = {
        "rating": "Rating 3.0",
        "kills": "击杀",
        "deaths": "死亡",
        "adr": "ADR",
        "headshot": "爆头率",
        "clutch": "残局胜率",
    }

    # ── 模式 1: 比赛详情 ──
    if match_id:
        url = f"https://www.hltv.org/matches/{match_id}"
        raw = _hltv_get(url, "match")
        if raw.startswith("ERROR:"):
            return f"错误: 无法获取比赛 {match_id} 的数据 ({raw})"
        soup = BeautifulSoup(raw, "html.parser")
        map_scores = _scrape_map_scores(soup)
        stats = _scrape_match_stats(soup, map_scores)
        if not stats:
            return f"错误: 无法解析比赛 {match_id} 的选手数据"
        output = f"HLTV 比赛选手 {stat_labels[stat_type]}\n{'=' * 50}\n"
        # 地图比分概览
        if map_scores:
            played = [ms for ms in map_scores if ms["status"] == "played"]
            if played:
                t1n = played[0]["team1"]["name"]
                t2n = played[0]["team2"]["name"]
                t1w = sum(1 for ms in played if ms["team1"]["result"] == "won")
                t2w = sum(1 for ms in played if ms["team2"]["result"] == "won")
                output += f"\n🏆 {t1n} [{t1w}] - [{t2w}] {t2n}\n"
                for ms in map_scores:
                    if ms["status"] == "played":
                        t1 = ms["team1"]
                        t2 = ms["team2"]
                        output += f"  {ms['map']}: {t1['name']} {t1['score']} - {t2['score']} {t2['name']}"
                        if ms["half_scores"]:
                            output += f"  {ms['half_scores']}"
                        output += "\n"
                output += "\n"
        for block in stats:
            output += f"\n📊 {block['map']} — {block.get('team', '')}\n{'-' * 30}\n"
            for p in block.get("players", []):
                output += (
                    f"  {p['name']:<28s}  K-D: {p['kd']:<7s}  "
                    f"ADR: {p['adr']:<6s}  Rating: {p['rating']}\n"
                )
        return output

    # ── 模式 2: 指定选手 ID ──
    if player_id:
        profile = _scrape_player_profile(player_id)
        if not profile:
            return f"错误: 无法获取选手 {player_id} 的数据"
        return _format_player_profile(profile, stat_type)

    # ── 模式 3: 选手名搜索 ──
    if query:
        players = _search_hltv_player(query)
        if not players:
            return f"未找到选手: {query}"

        # 精确匹配优先
        query_lower = query.lower()
        exact = [(n, p) for n, p in players if query_lower == n.lower() or query_lower in n.lower().split("'")[-1].strip("'")]
        candidates = exact if exact else players[:5]

        output = f"HLTV 搜索: {query}（找到 {len(players)} 个结果）\n{'=' * 50}\n"
        for name, path in candidates:
            profile = _scrape_player_profile(path)
            if profile:
                output += _format_player_profile(profile, stat_type)
                output += "\n" + "-" * 40 + "\n"
            else:
                output += f"  {name} — 无法获取详细数据\n"
        return output

    # ── 模式 4: 无参数（返回提示） ──
    return (
        "⚠️ HLTV /stats/players 端点受 Cloudflare 严格保护，此模式暂不可用。\n\n"
        "请使用以下方式获取选手数据:\n"
        "  1. 名称搜索: csgo_player_stats(query=\"Jee\")\n"
        "  2. 指定选手: csgo_player_stats(player_id=\"20702/jee\")\n"
        "  3. 比赛详情: csgo_player_stats(match_id=\"2394896/legacy-vs-tyloo-...\")\n"
        "  4. 5EPlay:   csgo_player_stats(query=\"Jee\", source=\"5eplay\")\n"
    )


def _format_player_profile(profile: dict, stat_type: str) -> str:
    """将 _scrape_player_profile 的结果格式化为文本输出。"""
    lines = []
    name = profile.get("name", "未知")
    team = profile.get("team", "")
    rating = profile.get("rating", "")
    stats = profile.get("stats", {})
    recent = profile.get("recent_matches", [])

    header = f"🎮 {name}"
    if team:
        header += f" | {team}"
    if rating:
        header += f" | Rating 3.0: {rating}"
    lines.append(header)

    # 详细属性
    if stats:
        other = {k: v for k, v in stats.items() if "Rating" not in k}
        if other:
            parts = [f"{k}: {v}" for k, v in list(other.items())[:5]]
            lines.append("  " + "  |  ".join(parts))

    # 近期比赛
    if recent:
        lines.append(f"\n  📅 近期比赛（共 {len(recent)} 场）:")
        for match in recent[:5]:
            texts = match if isinstance(match, list) else [match]
            line = "    " + "  |  ".join(str(t) for t in texts[:5])
            lines.append(line)

    return "\n".join(lines)


def _csgo_match_detail_5eplay(match_id: str) -> str:
    """从 5EPlay API 获取比赛详情（地图比分、veto、选手统计）。"""
    # 兼容多种 ID 格式: "csgo_mc_2394896" / "2394896" / "2394896/slug"
    mid = match_id.strip()
    if not mid.startswith("csgo_mc_"):
        # 从 HLTV 格式提取数字 ID
        m = re.match(r"(\d+)", mid)
        if m:
            mid = f"csgo_mc_{m.group(1)}"
        else:
            return f"错误: match_id 格式无效，例: csgo_mc_2394896 或 2394896"

    detail = _5eplay_match_detail(mid)
    if not detail:
        return f"错误: 无法获取 5EPlay 比赛 {mid} 的数据"

    mc = detail.get("mc_info", {})
    tt = detail.get("tt_info", {})
    gs = detail.get("global_state", {})
    t1_info = mc.get("t1_info", {})
    t2_info = mc.get("t2_info", {})

    t1n = t1_info.get("disp_name", "T1")
    t2n = t2_info.get("disp_name", "T2")
    t1s = gs.get("t1_score", "?")
    t2s = gs.get("t2_score", "?")
    fmt = mc.get("format", "?")

    # ═══ 输出 ═══
    output = f"5EPlay 比赛详情\n{'=' * 60}\n"

    # ── 概要 ──
    output += f"\n⚔️  {t1n}  vs  {t2n}\n"
    output += f"📡 赛事: {tt.get('disp_name', '?')}"
    output += f"  |  {tt.get('grade_label', '?')}"
    output += f"  |  奖金: {tt.get('bonus', '?')}"
    output += f"  |  地点: {tt.get('addr', '?')}\n"
    output += f"📊 总比分: {t1n} [{t1s}] — [{t2s}] {t2n}"
    output += f"  (BO{fmt})\n"

    # ── Veto ──
    bp_maps = gs.get("bp_map_item", [])
    if bp_maps:
        output += f"\n📋 地图选择 (Veto)\n{'-' * 40}\n"
        for bp in bp_maps:
            act = "❌ ban" if bp.get("bp_type") == "ban" else "✅ pick"
            side = t1n if bp.get("team_side") == "t1" else t2n
            output += f"  {side} {act} {bp.get('map_name', '?')}\n"
        # 标注 decider
        if len(bp_maps) >= 7:
            last = bp_maps[-1]
            output += f"  ⚡ {last.get('map_name', '?')} 被留下\n"

    # ── 每图比分 ──
    bouts = detail.get("bouts_state", [])
    played_bouts = [
        b for b in bouts
        if b.get("status") == "2" and b.get("t1_stats", {}).get("all_score", "")
    ]
    if played_bouts:
        output += f"\n🗺️  地图比分\n{'-' * 40}\n"
        for b in bouts:
            t1s_map = b.get("t1_stats", {})
            t2s_map = b.get("t2_stats", {})
            all1 = t1s_map.get("all_score", "")
            all2 = t2s_map.get("all_score", "")
            map_name = b.get("map_name", "?")

            if not all1 and not all2:
                output += f"  {map_name:<12s} (未进行)\n"
                continue

            fh1, fh2 = t1s_map.get("fh_score", "?"), t2s_map.get("fh_score", "?")
            sh1, sh2 = t1s_map.get("sh_score", "?"), t2s_map.get("sh_score", "?")
            ot1, ot2 = t1s_map.get("ot_score", "0"), t2s_map.get("ot_score", "0")

            half_str = f"(FH {fh1}:{fh2} SH {sh1}:{sh2}"
            if ot1 or ot2:
                half_str += f" OT {ot1}:{ot2}"
            half_str += ")"

            output += f"  {map_name:<12s} {t1n} {all1} — {all2} {t2n}"
            output += f"    {half_str}\n"

    # ── MVP + 赔率 ──
    mvp = gs.get("mvp_player_stats", {})
    if mvp and mvp.get("name"):
        output += f"\n🌟 MVP: {mvp.get('name', '?')}"
        output += f"  Rating: {mvp.get('rating', '?')}  ADR: {mvp.get('adr', '?')}"
        output += f"  K-D: {mvp.get('kill','?')}-{mvp.get('death','?')}\n"

    t1_odds = gs.get("t1_odds", "")
    if t1_odds:
        output += f"\n💰 赔率: {t1n} {t1_odds} ({gs.get('t1_odds_percent','?')}%)"
        output += f"  |  {t2n} {gs.get('t2_odds','?')} ({gs.get('t2_odds_percent','?')}%)\n"

    # ── 全场聚合统计 ──
    output += f"\n📊 全场选手统计\n{'=' * 60}\n"
    for side_key, side_name in [("t1_player_stats", t1n), ("t2_player_stats", t2n)]:
        players = gs.get(side_key, [])
        if not players:
            continue
        output += f"\n  {side_name}\n  {'-' * 40}\n"
        for p in players:
            name = p.get("name", "?")
            kd = f"{p.get('kill','?')}-{p.get('death','?')}"
            output += (
                f"  {name:<20s} "
                f"K-D: {kd:<8s} "
                f"ADR: {str(p.get('adr','?')):<7s} "
                f"KAST: {str(p.get('kast','?')):<7s} "
                f"HS%: {str(p.get('head_shot_rate','?')):<7s} "
                f"Rating: {p.get('rating','?')}\n"
            )

    # ── 每图选手统计 ──
    if played_bouts:
        output += f"\n📊 每图选手统计\n{'=' * 60}\n"
        for b in played_bouts:
            map_name = b.get("map_name", "?")
            output += f"\n📊 {map_name}\n{'-' * 40}\n"
            for side_key, side_name in [("t1_pr_stats", t1n), ("t2_pr_stats", t2n)]:
                players = b.get(side_key, [])
                if not players:
                    continue
                output += f"\n  {side_name}\n"
                for p in players:
                    name = p.get("name", "?")
                    kd = f"{p.get('kill','?')}-{p.get('death','?')}"
                    output += (
                        f"  {name:<20s} "
                        f"K-D: {kd:<8s} "
                        f"ADR: {str(p.get('adr','?')):<7s} "
                        f"KAST: {str(p.get('kast','?')):<7s} "
                        f"HS%: {str(p.get('head_shot_rate','?')):<7s} "
                        f"Rating: {p.get('rating','?')}\n"
                    )

    return output


def _csgo_player_stats_5eplay_match(match_id: str, stat_type: str) -> str:
    """从 5EPlay 比赛详情 API 获取指定比赛的选手统计。"""
    detail = _5eplay_match_detail(match_id)
    if not detail:
        return f"错误: 无法获取 5EPlay 比赛 {match_id} 的数据"

    mc = detail.get("mc_info", {})
    tt = detail.get("tt_info", {})
    gs = detail.get("global_state", {})
    t1_info = mc.get("t1_info", {})
    t2_info = mc.get("t2_info", {})

    t1n = t1_info.get("disp_name", "T1")
    t2n = t2_info.get("disp_name", "T2")
    t1s = gs.get("t1_score", "?")
    t2s = gs.get("t2_score", "?")

    output = f"5EPlay 比赛选手统计\n{'=' * 60}\n"
    output += f"\n⚔️  {t1n} [{t1s}] — [{t2s}] {t2n}\n"
    output += f"📡 {tt.get('disp_name', '?')}"
    output += f"  |  {tt.get('grade_label', '?')}"
    output += f"  |  奖金: {tt.get('bonus', '?')}"
    output += f"  |  地点: {tt.get('addr', '?')}\n"

    # Veto
    bp_maps = gs.get("bp_map_item", [])
    if bp_maps:
        output += f"\n📋 地图选择\n{'-' * 40}\n"
        for bp in bp_maps:
            act = "ban" if bp.get("bp_type") == "ban" else "pick"
            side = t1n if bp.get("team_side") == "t1" else t2n
            output += f"  {side} {act} {bp.get('map_name', '?')}\n"

    # 每图比分 + 选手统计
    bouts = detail.get("bouts_state", [])
    for b in bouts:
        if b.get("status") != "2":
            continue
        map_name = b.get("map_name", "?")
        t1s_map = b.get("t1_stats", {})
        t2s_map = b.get("t2_stats", {})

        output += f"\n🗺️  {map_name}: {t1s_map.get('all_score','?')} — {t2s_map.get('all_score','?')}"
        output += f"  (FH {t1s_map.get('fh_score','?')}:{t2s_map.get('fh_score','?')}"
        output += f" SH {t1s_map.get('sh_score','?')}:{t2s_map.get('sh_score','?')}"
        ot1 = t1s_map.get("ot_score", "0")
        ot2 = t2s_map.get("ot_score", "0")
        if ot1 or ot2:
            output += f" OT {ot1}:{ot2}"
        output += ")\n"

        # 选手统计
        for side_key, side_name in [("t1_pr_stats", t1n), ("t2_pr_stats", t2n)]:
            players = b.get(side_key, [])
            output += f"\n  {side_name}\n  {'-' * 30}\n"
            for p in players:
                name = p.get("name", "?")
                rating = p.get("rating", "?")
                adr = p.get("adr", "?")
                kd = f"{p.get('kill','?')}-{p.get('death','?')}"
                kast = p.get("kast", "?")
                hs = p.get("head_shot_rate", "?")
                output += (
                    f"  {name:<20s} "
                    f"K-D: {kd:<8s} "
                    f"ADR: {str(adr):<7s} "
                    f"KAST: {str(kast):<7s} "
                    f"HS%: {str(hs):<7s} "
                    f"Rating: {rating}\n"
                )

    return output


def _csgo_player_stats_5eplay(query: str, stat_type: str) -> str:
    """从 5EPlay API 获取选手统计（先搜 session_list，再按需拉 detail）。"""
    if not query:
        return (
            "5EPlay 选手统计需要提供 match_id 或 query 参数。\n"
            "用法: csgo_player_stats(query=\"队伍名\", source=\"5eplay\") 或\n"
            "      csgo_player_stats(match_id=\"csgo_mc_xxxxx\", source=\"5eplay\")"
        )

    # 第一阶段：只从 session_list 搜索匹配的比赛 ID（轻量，最多 2 页）
    candidate_ids: list[str] = []
    q = query.lower()
    for gs, pages in [(1, 1), (2, 2)]:  # 进行中1页，已结束2页
        for page in range(1, pages + 1):
            data = _5eplay_api_get(_5EPLAY_LIST_URL, {
                "game_status": gs, "game_type": 1, "grades": "",
                "page": page, "limit": 30,
            }, ttl=120 if gs == 1 else 600)
            if not data or not data.get("success"):
                continue
            for m in data.get("data", {}).get("matches", []):
                info = m.get("mc_info", {})
                t1n = info.get("t1_info", {}).get("disp_name", "").lower()
                t2n = info.get("t2_info", {}).get("disp_name", "").lower()
                if q in t1n or q in t2n:
                    mid = info.get("id", "")
                    if mid and mid not in candidate_ids:
                        candidate_ids.append(mid)
            if len(candidate_ids) >= 3:
                break
        if len(candidate_ids) >= 3:
            break

    if not candidate_ids:
        return f"5EPlay 未找到与 '{query}' 相关的比赛。\n提示: 5EPlay 只保留约 10 页近期比赛，较早的比赛请用 HLTV。"

    # 第二阶段：只对前 2 个候选拉取完整 detail（之前每个都拉，现在最多 2 次）
    output = f"5EPlay 选手统计: {query}\n{'=' * 60}\n"
    shown = 0
    for mid in candidate_ids[:2]:
        detail = _5eplay_match_detail(mid)
        if not detail:
            continue
        shown += 1

        mc = detail.get("mc_info", {})
        tt = detail.get("tt_info", {})
        t1_info = mc.get("t1_info", {})
        t2_info = mc.get("t2_info", {})

        output += f"\n📡 {tt.get('disp_name', '?')}"
        output += f"  |  {t1_info.get('disp_name', '?')} vs {t2_info.get('disp_name', '?')}\n"

        # 全场聚合统计
        gs = detail.get("global_state", {})
        for side_key, side_label in [("t1_player_stats", t1_info.get("disp_name", "T1")),
                                      ("t2_player_stats", t2_info.get("disp_name", "T2"))]:
            players = gs.get(side_key, [])
            if players:
                output += f"\n  {side_label}\n  {'-' * 30}\n"
                for p in players[:5]:
                    name = p.get("name", "?")
                    kd = f"{p.get('kill','?')}-{p.get('death','?')}"
                    output += (
                        f"  {name:<20s} "
                        f"K-D: {kd:<8s} "
                        f"ADR: {str(p.get('adr','?')):<7s} "
                        f"KAST: {str(p.get('kast','?')):<7s} "
                        f"Rating: {p.get('rating','?')}\n"
                    )

    if shown == 0:
        return f"5EPlay 找到匹配比赛但无法获取详情: {', '.join(candidate_ids[:3])}"
    if len(candidate_ids) > 2:
        output += f"\n... 还有 {len(candidate_ids) - 2} 场匹配，缩小搜索范围可获取更多"

    return output


# ══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    mcp.run(transport="streamable-http")
